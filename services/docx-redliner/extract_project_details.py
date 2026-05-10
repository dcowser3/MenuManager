#!/usr/bin/env python3
"""
Extract project details and menu content from a populated RSH menu template DOCX.
Reverses the pattern in generate_from_form.py — reads template table fields
and extracts menu content after the boundary marker.

Usage:
    python extract_project_details.py <docx_path>
"""

import sys
import json
import re
from docx import Document

ALLERGEN_KEYWORDS = [
    "allergen",
    "gluten",
    "dairy",
    "fish",
    "nuts",
    "egg",
    "vegan",
    "vegetarian",
    "crustacean",
    "soy",
    "sesame",
]


def _normalize_allergen_label(label: str) -> str:
    return " ".join(label.split()).strip().lower()


def _parse_parenthesized_allergen_key(text: str) -> str:
    """
    Parse legends formatted as "(C) CELERY (D) DAIRY ..." and normalize them to
    the pipe-delimited format used by the dashboard allergen-key field.
    """
    if not text or "(" not in text or ")" not in text:
        return ""

    # Footer copy often follows the legend in the same text extraction stream.
    # Trim it so the last allergen label does not absorb boilerplate text.
    text = re.split(
        r"\b(?:ALL\s+PRICES|WE\s+WELCOME|CONSUMPTION\s+OF\s+RAW|CONSUMING\s+RAW|FOODBORNE\s+ILLNESS)\b",
        text,
        flags=re.IGNORECASE,
        maxsplit=1,
    )[0]

    pairs = []
    pattern = re.compile(
        r"\(\s*([A-Za-z]{1,3})\s*\)\s*([A-Za-z][A-Za-z\s/&\-]*?)(?=\s*\(\s*[A-Za-z]{1,3}\s*\)|$)"
    )
    for match in pattern.finditer(text):
        code = match.group(1).upper()
        label = _normalize_allergen_label(match.group(2))
        if not label:
            continue
        pairs.append((code, label))

    if len(pairs) < 4:
        return ""

    keyword_hits = sum(
        1 for _, label in pairs
        if any(keyword in label for keyword in ALLERGEN_KEYWORDS)
    )
    if keyword_hits < 2:
        return ""

    return " | ".join(f"{code} {label}" for code, label in pairs)


def detect_allergen_key(paragraphs) -> str:
    """
    Best-effort extraction of allergen legend text from DOCX content.
    Returns a single line string like:
    "C crustaceans | D dairy | E egg | F fish | G gluten | N nuts"
    """
    candidate_lines = []
    for paragraph in paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            candidate_lines.append(" ".join(text.split()))

    for line in candidate_lines:
        lower = line.lower()

        if "allergen" in lower and "|" in line:
            return line

        if "|" not in line:
            continue

        parts = [p.strip() for p in line.split("|") if p.strip()]
        parsed = []
        for part in parts:
            match = re.match(r"^\*?\s*([A-Za-z]{1,3})\s+([A-Za-z][A-Za-z\s/&\-]{2,})$", part)
            if not match:
                continue
            code = match.group(1).upper()
            label = " ".join(match.group(2).split())
            parsed.append((code, label))

        if len(parsed) >= 4:
            keyword_hits = sum(
                1 for _, label in parsed
                if any(keyword in label.lower() for keyword in ALLERGEN_KEYWORDS)
            )
            if keyword_hits >= 2:
                return " | ".join(f"{code} {label}" for code, label in parsed)

    combined_text = " ".join(candidate_lines)
    parenthesized_key = _parse_parenthesized_allergen_key(combined_text)
    if parenthesized_key:
        return parenthesized_key

    return ""


def extract_project_details(docx_path: str) -> dict:
    """
    Extract project details from the template table and menu content
    from after the boundary marker.

    Returns:
        dict with project_details and menu_content
    """
    doc = Document(docx_path)

    # Extract project details from first table.
    # `property` holds the legacy single-field value (old template's PROPERTY row).
    # `outlet`/`hotel`/`city` are filled from the newer template that splits the
    # property identity across rows (OUTLET NAME / HOTEL NAME / CITY / COUNTRY).
    # The dashboard combines outlet+hotel+city to find the canonical catalog entry.
    project_details = {
        "project_name": "",
        "property": "",
        "outlet": "",
        "hotel": "",
        "city": "",
        "size": "",
        "orientation": "",
        "date_needed": ""
    }

    # Exact-match mapping for standard RSH template fields
    field_mapping = {
        "PROJECT NAME": "project_name",
        "MENU NAME": "project_name",
        "PROPERTY": "property",
        "OUTLET NAME": "outlet",
        "HOTEL NAME": "hotel",
        "CITY / COUNTRY": "city",
        "CITY/COUNTRY": "city",
        "SIZE (PIXELS = WEB) OR (INCHES = PRINT)": "size",
        "ORIENTATION (PORTRAIT OR LANDSCAPE)": "orientation",
        "DATE NEEDED": "date_needed"
    }

    # Normalized keyword mapping for briefs that use alternative field names.
    # Order matters; first match wins. Outlet/hotel/city are checked before
    # the generic "property" fallback so newer split templates extract them
    # cleanly instead of collapsing into `property`.
    alt_field_mapping = [
        # project_name alternatives
        ("project name", "project_name"),
        ("menu name", "project_name"),
        ("event name", "project_name"),
        # split-property fields (newer template)
        ("outlet name", "outlet"),
        ("outlet", "outlet"),
        ("hotel name", "hotel"),
        ("hotel", "hotel"),
        ("city / country", "city"),
        ("city/country", "city"),
        ("city", "city"),
        ("country", "city"),
        # legacy single-field property
        ("property", "property"),
        ("venue name", "property"),
        ("restaurant name", "property"),
        ("location name", "property"),
        ("location", "property"),
        # size
        ("size", "size"),
        ("dimension", "size"),
        # orientation
        ("orientation", "orientation"),
        # date
        ("date needed", "date_needed"),
        ("due date", "date_needed"),
        ("deadline", "date_needed"),
    ]

    if len(doc.tables) > 0:
        table = doc.tables[0]
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                field_name = cells[0].text.strip()
                value = cells[1].text.strip()
                if not value:
                    continue

                # Try exact match first (standard template)
                if field_name in field_mapping:
                    key = field_mapping[field_name]
                    if not project_details[key]:
                        project_details[key] = value
                    continue

                # Fuzzy match for alternative field names (briefs)
                field_lower = field_name.lower()
                for pattern, key in alt_field_mapping:
                    if pattern in field_lower and not project_details[key]:
                        project_details[key] = value
                        break

    # Find boundary marker and extract menu content after it
    boundary_markers = [
        "Please drop the menu content below on page 2",  # Food template
        "MENU"  # Beverage template
    ]
    boundary_index = None

    for i, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if para_text == "MENU":
            boundary_index = i
            break
        elif boundary_markers[0] in para_text:
            boundary_index = i
            break

    # Extract menu content after boundary
    # Also skip instruction lines that follow the boundary
    menu_lines = []
    if boundary_index is not None:
        for i in range(boundary_index + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text
            stripped = text.strip()
            # Skip the instruction line and standalone "MENU" markers
            if stripped == "MENU" or "Please drop the menu content below" in stripped:
                continue
            menu_lines.append(text)
    else:
        # No boundary found — try to get all text after the table
        # Skip paragraphs that look like header/template content
        in_content = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if in_content:
                menu_lines.append(paragraph.text)
            elif text and text not in field_mapping and not any(m in text for m in boundary_markers):
                # Heuristic: start capturing after we pass table-related content
                in_content = True
                menu_lines.append(paragraph.text)

    # Clean up: strip trailing empty lines
    while menu_lines and not menu_lines[-1].strip():
        menu_lines.pop()

    menu_content = "\n".join(menu_lines)
    allergen_key = detect_allergen_key(doc.paragraphs)

    return {
        "project_details": project_details,
        "menu_content": menu_content,
        "allergen_key": allergen_key
    }


def main():
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Usage: python extract_project_details.py <docx_path>"}))
        sys.exit(1)

    docx_path = sys.argv[1]

    try:
        result = extract_project_details(docx_path)
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
