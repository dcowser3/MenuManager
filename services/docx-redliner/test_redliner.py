#!/usr/bin/env python3
"""
Test Script for Menu Redliner
==============================
Creates a sample document with menu items and tests the redlining functionality.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from menu_redliner import MenuRedliner


def create_test_document(output_path: str = "test_menu.docx"):
    """
    Create a sample menu document with the template structure
    and deliberate errors for testing.
    """
    doc = Document()
    
    # --- PAGE 1: Template Section ---
    
    # Add title
    title = doc.add_heading('RSH Design Brief - Menu Review', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some template content
    doc.add_paragraph("Project Name: Test Restaurant")
    doc.add_paragraph("Date: October 2025")
    doc.add_paragraph("Reviewer: AI System")
    
    doc.add_paragraph()  # Spacing
    
    # Add styled template paragraph
    template_para = doc.add_paragraph()
    run1 = template_para.add_run("Important: ")
    run1.bold = True
    run1.font.size = Pt(12)
    run2 = template_para.add_run("This template section should remain unchanged.")
    run2.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Add the boundary marker
    boundary = doc.add_paragraph("Please drop the menu content below on page 2.")
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boundary_run = boundary.runs[0]
    boundary_run.bold = True
    boundary_run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # --- PAGE 2: Menu Content (with deliberate errors) ---
    
    # Menu header
    menu_header = doc.add_heading('Menu Items', level=2)
    
    # Test items with various types of errors and formatting
    test_items = [
        {
            'text': "Guacamole - Fresh avacado, lime, cilantro - $12",
            'formatting': [
                {'range': (0, 9), 'bold': True},  # "Guacamole" bold
            ]
        },
        {
            'text': "Ceasar Salad - Romaine lettuce, parmesian cheese, croutons - $14",
            'formatting': [
                {'range': (0, 12), 'bold': True},  # "Caesar Salad" bold
            ]
        },
        {
            'text': "Margherita Pizza - tomato sauce, mozarella, fresh basil - $16",
            'formatting': [
                {'range': (0, 16), 'bold': True, 'italic': True},  # "Margherita Pizza" bold+italic
            ]
        },
        {
            'text': "Grilled Salmon - Wild caught samon, lemon butter, asparagus - $24",
            'formatting': [
                {'range': (0, 14), 'bold': True},  # "Grilled Salmon" bold
            ]
        },
        {
            'text': "Chocolate Cake - Rich chocolote cake with rasberry coulis - $10",
            'formatting': [
                {'range': (0, 14), 'bold': True},
                {'range': (56, 64), 'italic': True},  # Price italic
            ]
        },
    ]
    
    for item_data in test_items:
        para = doc.add_paragraph()
        text = item_data['text']
        
        # If no formatting specified, just add the text
        if not item_data.get('formatting'):
            para.add_run(text)
            continue
        
        # Apply character-level formatting
        formatted_ranges = item_data['formatting']
        last_end = 0
        
        for fmt in formatted_ranges:
            start, end = fmt['range']
            
            # Add any text before this formatted range
            if start > last_end:
                para.add_run(text[last_end:start])
            
            # Add the formatted text
            run = para.add_run(text[start:end])
            if fmt.get('bold'):
                run.bold = True
            if fmt.get('italic'):
                run.italic = True
            if fmt.get('font_size'):
                run.font.size = Pt(fmt['font_size'])
            if fmt.get('color'):
                r, g, b = fmt['color']
                run.font.color.rgb = RGBColor(r, g, b)
            
            last_end = end
        
        # Add any remaining text
        if last_end < len(text):
            para.add_run(text[last_end:])
    
    # Add a section header to test structure preservation
    doc.add_paragraph()
    section_header = doc.add_heading('Beverages', level=3)
    
    # More test items
    beverage_items = [
        "Espresso - Double shot espreso - $4",
        "Green Tea - Organic green tee - $3",
    ]
    
    for text in beverage_items:
        para = doc.add_paragraph(text)
    
    # Save the document
    doc.save(output_path)
    print(f"✓ Created test document: {output_path}")
    return output_path


def simple_test_correction(text: str) -> str:
    """
    Simple correction function for testing without AI.
    Fixes common deliberate errors in test document.
    """
    corrections = {
        'avacado': 'avocado',
        'Ceasar': 'Caesar',
        'parmesian': 'parmesan',
        'mozarella': 'mozzarella',
        'samon': 'salmon',
        'chocolote': 'chocolate',
        'rasberry': 'raspberry',
        'espreso': 'espresso',
        'tee': 'tea',
    }
    
    corrected = text
    for wrong, right in corrections.items():
        corrected = corrected.replace(wrong, right)
    
    return corrected


def print_document_analysis(doc_path: str, title: str):
    """
    Print analysis of a document's content.
    """
    doc = Document(doc_path)
    print(f"\n{title}")
    print("=" * 60)
    
    marker_found = False
    menu_paras = []
    
    for i, para in enumerate(doc.paragraphs):
        if "Please drop the menu content below" in para.text:
            marker_found = True
            print(f"Boundary marker found at paragraph {i}")
            continue
        
        if marker_found and para.text.strip():
            menu_paras.append(para.text)
    
    print(f"Found {len(menu_paras)} menu paragraphs after marker")
    print("\nMenu content:")
    print("-" * 60)
    for text in menu_paras[:5]:  # Show first 5
        print(f"  {text}")
    if len(menu_paras) > 5:
        print(f"  ... and {len(menu_paras) - 5} more")


def test_with_simple_corrections():
    """
    Test the redliner with simple predefined corrections (no AI needed).
    """
    print("\n" + "=" * 60)
    print("Testing Menu Redliner (No AI Required)")
    print("=" * 60)
    
    # Create test document
    test_doc = create_test_document("test_menu_input.docx")
    
    # Analyze original
    print_document_analysis(test_doc, "Original Document")
    
    # Process the document
    print("\n" + "-" * 60)
    print("Processing document...")
    print("-" * 60)
    
    redliner = MenuRedliner()
    output_path = redliner.process_document(
        test_doc,
        simple_test_correction,
        "test_menu_output.docx"
    )
    
    # Analyze result
    print_document_analysis(output_path, "Processed Document")
    
    print("\n" + "=" * 60)
    print("✓ Test completed successfully!")
    print("=" * 60)
    print(f"\nCheck these files:")
    print(f"  Input:  {test_doc}")
    print(f"  Output: {output_path}")
    print("\nOpen the output file in Word to see:")
    print("  • Red strikethrough on errors (avacado, Ceasar, etc.)")
    print("  • Yellow highlight on corrections (avocado, Caesar, etc.)")
    print("  • All original formatting preserved (bold, italic, etc.)")


def test_with_ai():
    """
    Test the redliner with actual AI corrections.
    """
    try:
        from ai_corrector import AICorrector
        
        print("\n" + "=" * 60)
        print("Testing Menu Redliner with AI")
        print("=" * 60)
        
        # Create test document
        test_doc = create_test_document("test_menu_ai_input.docx")
        
        # Initialize AI corrector
        print("\nInitializing AI corrector...")
        corrector = AICorrector()
        
        # Process the document
        print("Processing document with AI corrections...")
        redliner = MenuRedliner()
        output_path = redliner.process_document(
            test_doc,
            corrector.correct_text,
            "test_menu_ai_output.docx"
        )
        
        print("\n" + "=" * 60)
        print("✓ AI Test completed successfully!")
        print("=" * 60)
        print(f"\nCheck these files:")
        print(f"  Input:  {test_doc}")
        print(f"  Output: {output_path}")
        
    except ValueError as e:
        print(f"\n⚠ Skipping AI test: {e}")
        print("To enable AI testing, set OPENAI_API_KEY environment variable")


def main():
    """Main test runner."""
    print("\n" + "=" * 70)
    print(" Menu Redliner Test Suite")
    print("=" * 70)
    
    # Test 1: Simple corrections (no AI needed)
    test_with_simple_corrections()
    
    # Test 2: AI corrections (if API key available)
    if os.getenv("OPENAI_API_KEY"):
        print("\n")
        test_with_ai()
    else:
        print("\n" + "=" * 60)
        print("Skipping AI tests (no OPENAI_API_KEY found)")
        print("=" * 60)
        print("\nTo test with AI:")
        print("  export OPENAI_API_KEY='your-key-here'")
        print("  python test_redliner.py")
    
    print("\n" + "=" * 70)
    print("All tests completed!")
    print("=" * 70)


if __name__ == "__main__":
    main()

