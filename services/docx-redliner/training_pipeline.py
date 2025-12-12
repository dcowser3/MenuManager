"""
Training Pipeline for Menu Redliner
====================================
This module enables learning from historical document pairs (original + human-redlined)
to continuously improve the AI correction model.

Features:
- Batch ingestion of document pairs
- Pattern extraction from human corrections
- Rule generation and refinement
- Prompt optimization
- Training metrics and analytics
"""

import os
import json
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from docx import Document
from collections import defaultdict, Counter
import difflib
from datetime import datetime
import re

# Import dish database for storing learned dish information
try:
    from dish_allergen_db import (
        store_allergen_correction,
        learn_dish_from_correction,
        store_dish_terminology_correction,
        store_approved_dish,
        get_dish_corrections,
        extract_restaurant,
        extract_menu_date,
        extract_menu_type,
        get_statistics as get_dish_db_stats
    )
    DISH_DB_AVAILABLE = True
except ImportError:
    DISH_DB_AVAILABLE = False
    print("Note: dish_allergen_db not available, dish learning disabled")

# Import known correction pairs
# ADD NEW PAIRS TO: known_corrections.py (single location!)
try:
    from known_corrections import (
        KNOWN_PAIRS, KNOWN_ABBREVIATIONS, TERMINOLOGY_CORRECTIONS, CONTEXT_HINTS,
        is_known_pair, is_terminology_correction, get_context_hints
    )
    KNOWN_CORRECTIONS_AVAILABLE = True
except ImportError:
    KNOWN_PAIRS = set()
    KNOWN_ABBREVIATIONS = {}
    TERMINOLOGY_CORRECTIONS = {}
    CONTEXT_HINTS = {}
    KNOWN_CORRECTIONS_AVAILABLE = False
    print("Note: known_corrections.py not found, using inline pairs")
    
    def is_terminology_correction(x): return False
    def get_context_hints(x, y): return {}


class TrainingPairAnalyzer:
    """
    Analyzes pairs of original and human-redlined documents to extract
    correction patterns and rules.
    """
    
    def __init__(self):
        self.corrections = []
        self.patterns = defaultdict(list)
        self.rule_candidates = []
        
    def load_document_pair(self, original_path: str, redlined_path: str) -> Dict:
        """
        Load and analyze a pair of documents.
        
        Args:
            original_path: Path to original document
            redlined_path: Path to human-redlined document
            
        Returns:
            Dictionary containing analysis results
        """
        original_doc = Document(original_path)
        redlined_doc = Document(redlined_path)
        
        # Extract text paragraphs
        original_paras = [p.text for p in original_doc.paragraphs if p.text.strip()]
        redlined_paras = [p.text for p in redlined_doc.paragraphs if p.text.strip()]
        
        # Also extract formatting information
        original_formatted = self._extract_formatted_content(original_doc)
        redlined_formatted = self._extract_formatted_content(redlined_doc)
        
        pair_analysis = {
            'original_path': original_path,
            'redlined_path': redlined_path,
            'timestamp': datetime.now().isoformat(),
            'text_corrections': [],
            'formatting_corrections': [],
            'tracked_changes': [],
            'metadata': {
                'original_paras': len(original_paras),
                'redlined_paras': len(redlined_paras)
            }
        }
        
        # Build set of text that was already highlighted in the original document
        # (these are menu edits by the chef, not reviewer corrections)
        original_highlighted_text = self._get_highlighted_text(original_formatted)
        
        # FIRST: Try to extract tracked changes from redlined doc formatting
        # (strikethrough = deleted, highlight = added)
        tracked_corrections = self._extract_tracked_changes(
            redlined_formatted, 
            original_highlighted_text
        )
        if tracked_corrections:
            pair_analysis['text_corrections'] = tracked_corrections
            pair_analysis['tracked_changes'] = tracked_corrections
        else:
            # Fallback: Analyze text differences between docs
            text_diffs = self._analyze_text_differences(original_paras, redlined_paras)
            pair_analysis['text_corrections'] = text_diffs
        
        # Analyze formatting differences
        format_diffs = self._analyze_formatting_differences(
            original_formatted, 
            redlined_formatted
        )
        pair_analysis['formatting_corrections'] = format_diffs
        
        return pair_analysis
    
    def _is_likely_correction(self, original: str, corrected: str) -> bool:
        """
        Check if this looks like a systematic correction (vs a random menu change).
        
        Returns True for patterns that look like corrections:
        - Abbreviation expansions (bbq → barbeque sauce)
        - Spacing fixes (sea food → seafood)
        - Diacritic additions (jalapeno → jalapeño)
        - Similar character patterns (high edit similarity)
        - Known terminology pairs
        
        Returns False for patterns that look like menu changes:
        - Completely different words with no similarity
        - Different food categories
        """
        if not original or not corrected:
            return True  # Can't determine, let it through
        
        orig_lower = original.lower().strip()
        corr_lower = corrected.lower().strip()
        
        # 1. Spacing fix: removing/adding spaces
        if orig_lower.replace(' ', '') == corr_lower.replace(' ', ''):
            return True
        
        # 2. Diacritic addition: same base letters
        import unicodedata
        def remove_diacritics(s):
            return ''.join(c for c in unicodedata.normalize('NFD', s) 
                          if unicodedata.category(c) != 'Mn')
        if remove_diacritics(orig_lower) == remove_diacritics(corr_lower):
            return True
        
        # 3. Abbreviation expansion: original is much shorter and starts similarly
        if orig_lower and len(orig_lower) <= 5 and len(corr_lower) > len(orig_lower):
            # Check if it's an abbreviation (first letters match)
            if corr_lower.startswith(orig_lower[0]):
                return True
            # Check common abbreviations (defined in known_corrections.py)
            if orig_lower in KNOWN_ABBREVIATIONS:
                return True
        
        # 4. Known terminology pairs (same concept, different term)
        # These are defined in known_corrections.py - edit that file to add new pairs!
        if (orig_lower, corr_lower) in KNOWN_PAIRS:
            return True
        
        # 5. High character similarity (edit distance)
        # Using simple ratio: shared characters / total characters
        orig_set = set(orig_lower.replace(' ', ''))
        corr_set = set(corr_lower.replace(' ', ''))
        if orig_set and corr_set:
            shared = len(orig_set & corr_set)
            total = len(orig_set | corr_set)
            similarity = shared / total
            # If more than 60% character overlap, likely a correction
            if similarity > 0.6:
                return True
        
        # 6. One contains the other (partial match)
        if orig_lower in corr_lower or corr_lower in orig_lower:
            return True
        
        # 7. Single word changes are often terminology preferences
        orig_words = orig_lower.split()
        corr_words = corr_lower.split()
        if len(orig_words) == 1 and len(corr_words) == 1:
            # Single word to single word - often a valid correction
            return True
        
        # If none of the above, it might be a menu change
        # Return False to let the occurrence filter handle it
        return False
    
    def _is_price_change(self, original: str, corrected: str) -> bool:
        """
        Check if this is a price change (should be ignored).
        
        Price changes are when both original and corrected are purely numeric,
        or contain only numbers and common price symbols ($, €, etc.)
        """
        # Strip common price-related characters
        price_chars = '$€£¥,.| '
        orig_clean = original.strip()
        corr_clean = corrected.strip()
        
        for char in price_chars:
            orig_clean = orig_clean.replace(char, '')
            corr_clean = corr_clean.replace(char, '')
        
        # If both are purely numeric after cleaning, it's a price change
        if orig_clean.isdigit() and corr_clean.isdigit():
            return True
        
        # Also check if either is empty and the other is numeric (price removal/addition)
        if (orig_clean.isdigit() and not corr_clean) or (corr_clean.isdigit() and not orig_clean):
            return True
        
        return False
    
    
    def _get_highlighted_text(self, formatted_content: List[Dict]) -> set:
        """
        Extract all text that is highlighted (yellow) in a document.
        
        This is used to identify text that was already highlighted in the original
        document before review - these are menu edits, not corrections we want to learn.
        """
        highlighted = set()
        
        for para_info in formatted_content:
            for run in para_info.get('runs', []):
                highlight = run.get('highlight')
                if highlight and highlight != 'None' and 'YELLOW' in str(highlight).upper():
                    text = run.get('text', '').strip().lower()
                    if text:
                        highlighted.add(text)
        
        return highlighted
    
    def _extract_tracked_changes(self, formatted_content: List[Dict], original_highlighted_text: set = None) -> List[Dict]:
        """
        Extract corrections from tracked changes in a redlined document.
        
        Looks for patterns like:
        - [highlighted text][strikethrough text] = replacement (new before old)
        - [strikethrough text][highlighted text] = replacement (old before new)
        - [strikethrough text] alone = deletion
        - [highlighted text] alone = addition
        
        Filters out:
        - Price changes (numeric values)
        - Text that was already highlighted in the original (chef's menu edits)
        
        Note: Use --min-occurrences to filter out one-off menu changes vs systematic corrections
        
        Returns list of corrections found.
        """
        corrections = []
        original_highlighted_text = original_highlighted_text or set()
        
        for para_info in formatted_content:
            runs = para_info.get('runs', [])
            if not runs:
                continue
            
            # Track which runs we've already processed
            processed = set()
            
            # First pass: find highlight + strikethrough pairs
            for i, run in enumerate(runs):
                if i in processed:
                    continue
                    
                run_text = run.get('text', '').strip()
                if not run_text:
                    continue
                
                highlight = run.get('highlight')
                is_highlighted = highlight and highlight != 'None' and 'YELLOW' in str(highlight).upper()
                is_strikethrough = run.get('strike')
                
                # Pattern 1: Highlighted text followed by strikethrough (replacement: new → old)
                if is_highlighted:
                    added_text = run_text
                    deleted_text = ''
                    
                    # Look for following strikethrough
                    for j in range(i + 1, min(i + 4, len(runs))):  # Check next few runs
                        if j in processed:
                            continue
                        next_run = runs[j]
                        next_text = next_run.get('text', '').strip()
                        
                        if not next_text:
                            continue
                        
                        if next_run.get('strike'):
                            deleted_text = next_text
                            processed.add(j)
                            break
                        # Stop if we hit regular text
                        elif not (next_run.get('highlight') and 'YELLOW' in str(next_run.get('highlight', '')).upper()):
                            break
                    
                    if deleted_text:
                        # Skip price changes
                        if self._is_price_change(deleted_text, added_text):
                            processed.add(i)
                            continue
                        
                        # Skip if the "new" text was already highlighted in the original
                        # (this is a menu edit by the chef, not a reviewer correction)
                        if added_text.lower() in original_highlighted_text:
                            processed.add(i)
                            continue
                        
                        # This is a replacement: old text was deleted, new text was added
                        category = self._categorize_correction(deleted_text, added_text)
                        
                        # Capture paragraph context (contains dish name)
                        para_text = para_info.get('text', '')
                        dish_name = self._extract_dish_name_from_para(para_text)
                        
                        correction = {
                            'type': 'replacement',
                            'original': deleted_text,
                            'corrected': added_text,
                            'category': category,
                            'source': 'tracked_changes',
                            'context': {
                                'paragraph': para_text[:100],  # First 100 chars for context
                                'dish_name': dish_name,
                            },
                            'word_diffs': [{
                                'operation': 'replace',
                                'original_words': deleted_text,
                                'corrected_words': added_text
                            }]
                        }
                        corrections.append(correction)
                        processed.add(i)
                        print(f"    Found tracked change: '{deleted_text}' → '{added_text}'")
                
                # Pattern 2: Strikethrough text followed by highlighted (replacement: old → new)
                elif is_strikethrough:
                    deleted_text = run_text
                    added_text = ''
                    
                    # Look for following highlighted text
                    for j in range(i + 1, min(i + 4, len(runs))):
                        if j in processed:
                            continue
                        next_run = runs[j]
                        next_text = next_run.get('text', '').strip()
                        
                        if not next_text:
                            continue
                        
                        next_highlight = next_run.get('highlight')
                        if next_highlight and next_highlight != 'None' and 'YELLOW' in str(next_highlight).upper():
                            added_text = next_text
                            processed.add(j)
                            break
                        # Stop if we hit regular non-strikethrough text
                        elif not next_run.get('strike'):
                            break
                    
                    if i not in processed:  # Not already processed as part of a highlight pair
                        # Skip price changes
                        if self._is_price_change(deleted_text, added_text):
                            processed.add(i)
                            continue
                        
                        # Skip if the "new" text was already highlighted in the original
                        if added_text and added_text.lower() in original_highlighted_text:
                            processed.add(i)
                            continue
                        
                        if added_text:
                            category = self._categorize_correction(deleted_text, added_text)
                        else:
                            category = 'deletion'
                        
                        # Capture paragraph context (contains dish name)
                        para_text = para_info.get('text', '')
                        dish_name = self._extract_dish_name_from_para(para_text)
                        
                        correction = {
                            'type': 'replacement' if added_text else 'deletion',
                            'original': deleted_text,
                            'corrected': added_text,
                            'category': category,
                            'source': 'tracked_changes',
                            'context': {
                                'paragraph': para_text[:100],
                                'dish_name': dish_name,
                            },
                            'word_diffs': [{
                                'operation': 'replace' if added_text else 'delete',
                                'original_words': deleted_text,
                                'corrected_words': added_text
                            }]
                        }
                        corrections.append(correction)
                        processed.add(i)
                        print(f"    Found tracked change: '{deleted_text}' → '{added_text}'")
        
        return corrections
    
    def _extract_formatted_content(self, doc: Document) -> List[Dict]:
        """
        Extract text content with formatting information.
        """
        formatted_content = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            para_info = {
                'text': para.text,
                'alignment': str(para.alignment),
                'runs': []
            }
            
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'strike': run.font.strike,
                    'highlight': str(run.font.highlight_color) if run.font.highlight_color else None,
                    'color': str(run.font.color.rgb) if run.font.color.rgb else None
                }
                para_info['runs'].append(run_info)
            
            formatted_content.append(para_info)
        
        return formatted_content
    
    def _analyze_text_differences(
        self, 
        original: List[str], 
        redlined: List[str]
    ) -> List[Dict]:
        """
        Analyze text-level differences between documents.
        """
        corrections = []
        
        # Use difflib to find matching paragraphs
        matcher = difflib.SequenceMatcher(None, original, redlined)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                # Text was changed
                for orig_idx in range(i1, i2):
                    for redl_idx in range(j1, j2):
                        correction = self._extract_correction_details(
                            original[orig_idx],
                            redlined[redl_idx]
                        )
                        if correction:
                            corrections.append(correction)
            elif tag == 'delete':
                # Content was removed
                for orig_idx in range(i1, i2):
                    corrections.append({
                        'type': 'deletion',
                        'original': original[orig_idx],
                        'corrected': '',
                        'category': 'content_removal'
                    })
            elif tag == 'insert':
                # Content was added
                for redl_idx in range(j1, j2):
                    corrections.append({
                        'type': 'insertion',
                        'original': '',
                        'corrected': redlined[redl_idx],
                        'category': 'content_addition'
                    })
        
        return corrections
    
    def _extract_correction_details(self, original: str, corrected: str) -> Optional[Dict]:
        """
        Extract detailed information about a specific correction.
        """
        if original == corrected:
            return None
        
        # Get word-level diffs first (needed for allergen detection)
        word_diffs = self._get_word_level_diffs(original, corrected)
        
        # Check if any word diff is an allergen correction
        category = self._categorize_correction_with_diffs(original, corrected, word_diffs)
        
        correction = {
            'type': 'replacement',
            'original': original,
            'corrected': corrected,
            'category': category,
            'word_diffs': word_diffs
        }
        
        return correction
    
    def _categorize_correction_with_diffs(self, original: str, corrected: str, word_diffs: List[Dict]) -> str:
        """
        Categorize correction using both full text and word-level diffs.
        Word diffs help catch allergen codes that are part of larger lines.
        """
        # First check if this looks like a "swapped menu item" (completely different content)
        # These are NOT corrections - they're just items in different order between documents
        if self._is_swapped_menu_item(original, corrected):
            return 'swapped_item'  # Will be filtered out later
        
        # Check word-level diffs for allergen patterns
        for wd in word_diffs:
            orig_words = wd.get('original_words', '')
            corr_words = wd.get('corrected_words', '')
            if self._is_allergen_correction(orig_words, corr_words):
                return 'allergen'
        
        # Fall back to full-text categorization
        return self._categorize_correction(original, corrected)
    
    def _is_swapped_menu_item(self, original: str, corrected: str) -> bool:
        """
        Detect if this is a swapped menu item (completely different dish) rather than a correction.
        
        Signs of a swapped item:
        - Very different lengths
        - Very low word overlap
        - Both are long strings (full menu item descriptions)
        """
        # Short strings are likely actual corrections
        if len(original) < 20 or len(corrected) < 20:
            return False
        
        # Get words (ignore case and punctuation)
        orig_words = set(re.sub(r'[^\w\s]', '', original.lower()).split())
        corr_words = set(re.sub(r'[^\w\s]', '', corrected.lower()).split())
        
        # Remove common filler words
        filler = {'the', 'a', 'an', 'and', 'or', 'with', 'of', 'in', 'on', 'for'}
        orig_words -= filler
        corr_words -= filler
        
        if not orig_words or not corr_words:
            return False
        
        # Calculate overlap
        overlap = orig_words & corr_words
        total_words = orig_words | corr_words
        
        overlap_ratio = len(overlap) / len(total_words) if total_words else 0
        
        # If less than 30% word overlap on long strings, it's probably a swapped item
        if overlap_ratio < 0.3:
            return True
        
        return False
    
    def _categorize_correction(self, original: str, corrected: str) -> str:
        """
        Automatically categorize the type of correction.
        """
        # Check for common patterns
        orig_lower = original.lower().strip()
        corr_lower = corrected.lower().strip()
        
        # Allergen code correction (check FIRST - these look like spelling but aren't)
        if self._is_allergen_correction(original, corrected):
            return 'allergen'
        
        # Terminology/word preference (check BEFORE spelling)
        # These are word substitutions, not misspellings: mayo→aioli, crust→rim
        if self._is_terminology_preference(orig_lower, corr_lower):
            return 'terminology'
        
        # Spelling correction
        if self._is_spelling_correction(original, corrected):
            return 'spelling'
        
        # Case change
        if original.lower() == corrected.lower():
            return 'case_change'
        
        # Punctuation/separator change
        if self._is_punctuation_change(original, corrected):
            return 'punctuation'
        
        # Diacritic addition
        if self._is_diacritic_change(original, corrected):
            return 'diacritics'
        
        # Price format
        if '$' in original or '$' in corrected or '|' in original or '|' in corrected:
            return 'price_format'
        
        # Separator change (/ vs -)
        if (' / ' in corrected and ' - ' in original) or (' - ' in corrected and ' / ' in original):
            return 'separator'
        
        return 'general'
    
    def _is_terminology_preference(self, original: str, corrected: str) -> bool:
        """
        Check if this is a terminology/word preference correction.
        
        Terminology corrections are word substitutions where neither word is
        "wrong" per se, but RSH prefers one over the other:
        - mayo → aioli (RSH prefers "aioli")
        - crust → rim (for cocktail glasses)
        - sorbete → sorbet
        
        These are different from spelling corrections because:
        - Both words are valid English/Spanish
        - The "original" isn't misspelled
        - It's a style/preference choice
        """
        # Check if it's in our known terminology corrections
        if original in TERMINOLOGY_CORRECTIONS:
            return True
        
        # Check if it's a known pair that's terminology (not spelling)
        if (original, corrected) in KNOWN_PAIRS:
            # Pairs that are terminology preferences, not misspellings
            terminology_pairs = {
                ('mayo', 'aioli'), ('aioli', 'mayo'),
                ('crust', 'rim'),
                ('shrimp', 'prawn'), ('prawn', 'shrimp'),
                ('tartare', 'tartar'), ('tartar', 'tartare'),
            }
            if (original, corrected) in terminology_pairs:
                return True
        
        # Single word to single word with low character similarity
        # but both are real words = likely terminology preference
        if len(original.split()) == 1 and len(corrected.split()) == 1:
            # Calculate character similarity
            orig_set = set(original)
            corr_set = set(corrected)
            if orig_set and corr_set:
                shared = len(orig_set & corr_set)
                total = len(orig_set | corr_set)
                similarity = shared / total
                
                # Low similarity + both short = different words, not typo
                if similarity < 0.4 and len(original) <= 6 and len(corrected) <= 6:
                    return True
        
        return False
    
    def _is_allergen_correction(self, original: str, corrected: str) -> bool:
        """
        Check if this is an allergen/dietary code correction.
        
        Allergen codes are comma-separated letters at the end of menu items:
        - D = Dairy
        - N = Nuts  
        - G = Gluten
        - V = Vegetarian
        - S = Vegan (or S* for special)
        - E = Eggs
        - F = Fish
        - C = Crustaceans
        - SE = Sesame
        - SY = Soy
        - M = Mustard
        
        Examples:
        - "d,n" → "d,n,v" (missing vegetarian marker)
        - "s" → "s*" (adding asterisk for raw/undercooked)
        - "v" → "d,v" (missing dairy marker)
        """
        # Valid allergen codes (single letters and two-letter codes)
        allergen_codes = {'d', 'n', 'g', 'v', 's', 'e', 'f', 'c', 'm', 'se', 'sy', 's*'}
        
        # Pattern for allergen strings: comma-separated codes, possibly with asterisks
        allergen_pattern = r'^[a-z,\*]+$'
        
        orig_clean = original.lower().strip()
        corr_clean = corrected.lower().strip()
        
        # Check if both strings look like allergen codes
        if re.match(allergen_pattern, orig_clean) and re.match(allergen_pattern, corr_clean):
            # Split by comma and check if they're valid codes
            orig_codes = set(orig_clean.replace('*', '').split(','))
            corr_codes = set(corr_clean.replace('*', '').split(','))
            
            # Remove empty strings
            orig_codes.discard('')
            corr_codes.discard('')
            
            # If most codes are valid allergen codes, it's an allergen correction
            all_codes = orig_codes | corr_codes
            valid_count = sum(1 for code in all_codes if code in allergen_codes)
            
            if valid_count >= len(all_codes) * 0.5:  # At least half are valid codes
                return True
        
        # Also check for single letter changes that are allergen codes
        if len(orig_clean) <= 3 and len(corr_clean) <= 3:
            if orig_clean in allergen_codes or corr_clean in allergen_codes:
                return True
        
        return False
    
    def _is_spelling_correction(self, original: str, corrected: str) -> bool:
        """Check if this is a spelling correction."""
        # Simple heuristic: similar length, high character overlap
        if abs(len(original) - len(corrected)) > 3:
            return False
        
        orig_words = set(original.lower().split())
        corr_words = set(corrected.lower().split())
        
        # Check for word substitutions that might be spelling
        if len(orig_words) == len(corr_words):
            diff_words = orig_words.symmetric_difference(corr_words)
            if len(diff_words) <= 2:  # Only 1-2 words changed
                return True
        
        return False
    
    def _is_punctuation_change(self, original: str, corrected: str) -> bool:
        """Check if this is primarily a punctuation change."""
        # Remove all punctuation and compare
        orig_clean = re.sub(r'[^\w\s]', '', original)
        corr_clean = re.sub(r'[^\w\s]', '', corrected)
        
        return orig_clean == corr_clean
    
    def _is_diacritic_change(self, original: str, corrected: str) -> bool:
        """Check if diacritics were added/changed."""
        import unicodedata
        
        # Normalize to remove diacritics
        orig_norm = unicodedata.normalize('NFD', original)
        orig_base = ''.join(c for c in orig_norm if unicodedata.category(c) != 'Mn')
        
        corr_norm = unicodedata.normalize('NFD', corrected)
        corr_base = ''.join(c for c in corr_norm if unicodedata.category(c) != 'Mn')
        
        return orig_base.lower() == corr_base.lower()
    
    def _get_word_level_diffs(self, original: str, corrected: str) -> List[Dict]:
        """
        Get word-level differences between two strings.
        """
        orig_words = original.split()
        corr_words = corrected.split()
        
        matcher = difflib.SequenceMatcher(None, orig_words, corr_words)
        diffs = []
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag != 'equal':
                diffs.append({
                    'operation': tag,
                    'original_words': ' '.join(orig_words[i1:i2]),
                    'corrected_words': ' '.join(corr_words[j1:j2])
                })
        
        return diffs
    
    def _extract_dish_name_from_para(self, paragraph_text: str) -> Optional[str]:
        """
        Extract the dish/item name from a menu paragraph.
        
        Menu items typically follow formats like:
        - "Dish Name, ingredient, ingredient, allergens PRICE"
        - "Dish Name - description"
        - "Red Paloma, tequila, grapefruit, lime, salt rim 15"
        
        Returns the dish name or None if can't parse.
        """
        if not paragraph_text or not paragraph_text.strip():
            return None
        
        text = paragraph_text.strip()
        
        # Try comma-separated format first (most common)
        if ', ' in text:
            dish_name = text.split(', ')[0].strip()
            # Clean up any leading formatting markers
            dish_name = re.sub(r'^[\d\.\)\-\s]+', '', dish_name)
            if dish_name and len(dish_name) > 1:
                return dish_name
        
        # Try dash-separated format
        if ' - ' in text:
            dish_name = text.split(' - ')[0].strip()
            dish_name = re.sub(r'^[\d\.\)\-\s]+', '', dish_name)
            if dish_name and len(dish_name) > 1:
                return dish_name
        
        # Fall back to first few words (before any numbers or special chars)
        words = text.split()
        if words:
            # Take words until we hit a price or allergen code
            name_parts = []
            for word in words:
                # Stop at numbers (prices) or single-letter allergen codes
                if word.isdigit() or (len(word) <= 2 and word.isupper()):
                    break
                name_parts.append(word)
                if len(name_parts) >= 4:  # Max 4 words for dish name
                    break
            
            if name_parts:
                return ' '.join(name_parts)
        
        return None
    
    def _analyze_formatting_differences(
        self, 
        original: List[Dict], 
        redlined: List[Dict]
    ) -> List[Dict]:
        """
        Analyze formatting differences between documents.
        """
        format_corrections = []
        
        # Compare formatting for similar paragraphs
        for orig_para in original:
            for redl_para in redlined:
                # If text is similar, check formatting
                if self._text_similarity(orig_para['text'], redl_para['text']) > 0.8:
                    if orig_para['alignment'] != redl_para['alignment']:
                        format_corrections.append({
                            'type': 'alignment',
                            'text': orig_para['text'][:50],
                            'original': orig_para['alignment'],
                            'corrected': redl_para['alignment']
                        })
        
        return format_corrections
    
    def _text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity ratio between two texts."""
        return difflib.SequenceMatcher(None, text1, text2).ratio()


class RuleGenerator:
    """
    Generates rules from analyzed correction patterns.
    """
    
    def __init__(self):
        self.rule_templates = {
            'allergen': 'Missing allergen/dietary code: "{original}" should include "{corrected}" - verify dish ingredients match allergen markers',
            'terminology': 'Use preferred terminology: "{original}" → "{corrected}"',
            'spelling': 'Correct spelling: "{original}" → "{corrected}"',
            'diacritics': 'Use proper diacritics: "{original}" → "{corrected}"',
            'punctuation': 'Fix punctuation: "{original}" → "{corrected}"',
            'separator': 'Use correct separator: "{original}" → "{corrected}"',
            'case_change': 'Apply correct case: "{original}" → "{corrected}"',
            'price_format': 'Format price correctly: "{original}" → "{corrected}"'
        }
    
    def _is_likely_correction(self, original: str, corrected: str) -> bool:
        """
        Check if this looks like a systematic correction (vs a random menu change).
        
        Returns True for patterns that look like corrections:
        - Abbreviation expansions (bbq → barbeque sauce)
        - Spacing fixes (sea food → seafood)
        - Diacritic additions (jalapeno → jalapeño)
        - Similar character patterns
        - Known terminology pairs
        - Single word to single word changes
        """
        if not original or not corrected:
            return True  # Can't determine, let it through
        
        orig_lower = original.lower().strip()
        corr_lower = corrected.lower().strip()
        
        # 1. Spacing fix: removing/adding spaces
        if orig_lower.replace(' ', '') == corr_lower.replace(' ', ''):
            return True
        
        # 2. Diacritic addition: same base letters
        import unicodedata
        def remove_diacritics(s):
            return ''.join(c for c in unicodedata.normalize('NFD', s) 
                          if unicodedata.category(c) != 'Mn')
        if remove_diacritics(orig_lower) == remove_diacritics(corr_lower):
            return True
        
        # 3. Abbreviation expansion: original is short, corrected is longer
        if orig_lower and len(orig_lower) <= 5 and len(corr_lower) > len(orig_lower):
            if corr_lower.startswith(orig_lower[0]):
                return True
            # Check common abbreviations (defined in known_corrections.py)
            if orig_lower in KNOWN_ABBREVIATIONS:
                return True
        
        # 4. Known terminology pairs
        # These are defined in known_corrections.py - edit that file to add new pairs!
        if (orig_lower, corr_lower) in KNOWN_PAIRS:
            return True
        
        # 5. High character similarity
        orig_set = set(orig_lower.replace(' ', ''))
        corr_set = set(corr_lower.replace(' ', ''))
        if orig_set and corr_set:
            shared = len(orig_set & corr_set)
            total = len(orig_set | corr_set)
            if shared / total > 0.6:
                return True
        
        # 6. One contains the other
        if orig_lower in corr_lower or corr_lower in orig_lower:
            return True
        
        # 7. Single word to single word - often terminology preferences
        if len(orig_lower.split()) == 1 and len(corr_lower.split()) == 1:
            return True
        
        return False
    
    def generate_rules_from_corrections(
        self, 
        corrections: List[Dict],
        min_occurrences: int = 2,
        existing_rules: List[Dict] = None
    ) -> List[Dict]:
        """
        Generate rules from a list of corrections.
        
        Args:
            corrections: List of correction dictionaries
            min_occurrences: Minimum times a pattern must appear to become a rule
            existing_rules: Optional list of existing rules to avoid duplicates
            
        Returns:
            List of generated rules (only NEW rules, not duplicates)
        """
        # Categories to skip (these are not real corrections)
        skip_categories = {'swapped_item', 'content_removal', 'content_addition'}
        
        # Build set of existing patterns to avoid duplicates
        existing_patterns = set()
        if existing_rules:
            for rule in existing_rules:
                if rule.get('rule_id', '').startswith('LEARNED'):
                    orig = rule.get('details', {}).get('original_text', '').lower()
                    corr = rule.get('details', {}).get('corrected_text', '').lower()
                    existing_patterns.add((orig, corr))
        
        # Group corrections by category
        by_category = defaultdict(list)
        for corr in corrections:
            if corr.get('type') == 'replacement':
                category = corr.get('category', 'unknown')
                if category not in skip_categories:
                    by_category[category].append(corr)
        
        generated_rules = []
        
        # Find patterns within each category
        for category, corr_list in by_category.items():
            pattern_counts, pattern_contexts = self._find_patterns(corr_list)
            
            for pattern, occurrences in pattern_counts.items():
                orig, corr = pattern
                
                # Skip if this pattern already exists
                if pattern in existing_patterns:
                    continue
                
                # Skip patterns that look like swapped items (long strings with low similarity)
                if len(orig) > 30 and len(corr) > 30:
                    # Check word overlap
                    orig_words = set(orig.split())
                    corr_words = set(corr.split())
                    overlap = len(orig_words & corr_words) / max(len(orig_words | corr_words), 1)
                    if overlap < 0.3:
                        continue  # Skip - likely swapped item
                
                # Get contexts for this pattern
                contexts = pattern_contexts.get(pattern, [])
                
                # Include rule if:
                # 1. It meets the minimum occurrence threshold, OR
                # 2. It looks like a systematic correction (passes heuristics)
                if occurrences >= min_occurrences:
                    rule = self._create_rule(category, pattern, occurrences, contexts)
                    generated_rules.append(rule)
                elif self._is_likely_correction(orig, corr):
                    # Single occurrence but looks like a real correction
                    rule = self._create_rule(category, pattern, occurrences, contexts)
                    generated_rules.append(rule)
                # Otherwise skip - likely a one-off menu change
        
        return generated_rules
    
    def _find_patterns(self, corrections: List[Dict]) -> Tuple[Dict[Tuple[str, str], int], Dict[Tuple[str, str], List[Dict]]]:
        """
        Find recurring patterns in corrections.
        
        Returns:
            Tuple of (pattern_count, pattern_contexts)
            - pattern_count: Dict mapping pattern to occurrence count
            - pattern_contexts: Dict mapping pattern to list of context dicts
        """
        pattern_count = Counter()
        pattern_contexts = defaultdict(list)
        
        for corr in corrections:
            # Get context info
            context = corr.get('context', {})
            
            # Extract word-level changes
            for word_diff in corr.get('word_diffs', []):
                if word_diff['operation'] == 'replace':
                    orig = word_diff['original_words'].lower()
                    corr_text = word_diff['corrected_words'].lower()
                    
                    # Store the pattern and its context
                    pattern = (orig, corr_text)
                    pattern_count[pattern] += 1
                    if context:
                        pattern_contexts[pattern].append(context)
        
        return pattern_count, pattern_contexts
    
    def _create_rule(
        self, 
        category: str, 
        pattern: Tuple[str, str], 
        occurrences: int,
        contexts: List[Dict] = None
    ) -> Dict:
        """
        Create a rule from a pattern.
        
        Args:
            category: The correction category (terminology, spelling, etc.)
            pattern: Tuple of (original, corrected) text
            occurrences: Number of times this pattern was seen
            contexts: List of context dicts with dish_name, paragraph info
        """
        original, corrected = pattern
        contexts = contexts or []
        
        # Extract dish names from contexts
        dish_names = []
        for ctx in contexts:
            if ctx.get('dish_name'):
                dish_names.append(ctx['dish_name'])
        dish_names = list(set(dish_names))  # Unique names
        
        # Get context hints if available (from known_corrections.py)
        hints = get_context_hints(original, corrected) if KNOWN_CORRECTIONS_AVAILABLE else {}
        
        rule = {
            'rule_id': f'LEARNED-{category.upper()}-{abs(hash(pattern)) % 10000:04d}',
            'category': category.replace('_', ' ').title(),
            'severity': 'Medium',
            'description': self.rule_templates.get(category, 'Fix: "{original}" → "{corrected}"').format(
                original=original,
                corrected=corrected
            ),
            'details': {
                'pattern_type': category,
                'original_text': original,
                'corrected_text': corrected,
                'occurrences': occurrences,
                'learned_from': 'training_data',
                'confidence': min(occurrences / 10.0, 1.0)  # Confidence based on occurrences
            }
        }
        
        # Add context information if available
        if dish_names:
            rule['details']['seen_on_dishes'] = dish_names
            # Increase confidence if we have dish context
            rule['details']['confidence'] = min(rule['details']['confidence'] + 0.1, 1.0)
        
        # Add context hints (item types, keywords) if available
        if hints:
            rule['details']['applies_to'] = hints.get('item_types', [])
            rule['details']['context_keywords'] = hints.get('keywords', [])
            rule['details']['note'] = hints.get('note', '')
        
        return rule


class PromptOptimizer:
    """
    Optimizes the AI system prompt based on learned corrections.
    """
    
    def __init__(self, current_prompt: str):
        self.current_prompt = current_prompt
        self.examples = []
    
    def add_training_examples(self, corrections: List[Dict]):
        """
        Add training examples from corrections.
        """
        for corr in corrections:
            if corr.get('type') == 'replacement':
                self.examples.append({
                    'input': corr['original'],
                    'output': corr['corrected'],
                    'category': corr.get('category', 'general')
                })
    
    def generate_enhanced_prompt(self, max_examples: int = 10) -> str:
        """
        Generate an enhanced prompt with learned examples.
        """
        # Count examples by category
        by_category = defaultdict(list)
        for ex in self.examples:
            by_category[ex['category']].append(ex)
        
        # Build examples section
        examples_text = "\n\nLEARNED EXAMPLES FROM TRAINING DATA:\n"
        
        example_count = 0
        for category in ['allergen', 'spelling', 'diacritics', 'separator', 'punctuation', 'price_format']:
            if category in by_category and example_count < max_examples:
                examples_text += f"\n{category.replace('_', ' ').title()}:\n"
                
                # Add up to 2 examples per category
                for ex in by_category[category][:2]:
                    examples_text += f'Input: "{ex["input"]}"\n'
                    examples_text += f'Output: "{ex["output"]}"\n\n'
                    example_count += 1
                    
                    if example_count >= max_examples:
                        break
        
        # Insert examples before the existing examples section
        enhanced_prompt = self.current_prompt
        if "EXAMPLES:" in enhanced_prompt:
            enhanced_prompt = enhanced_prompt.replace("EXAMPLES:", examples_text + "\n\nEXAMPLES:")
        else:
            enhanced_prompt += examples_text
        
        return enhanced_prompt


class TrainingPipeline:
    """
    Main training pipeline coordinator.
    """
    
    def __init__(self, training_data_dir: str = "tmp/training"):
        self.training_data_dir = Path(training_data_dir)
        self.training_data_dir.mkdir(parents=True, exist_ok=True)
        
        self.analyzer = TrainingPairAnalyzer()
        self.rule_generator = RuleGenerator()
        
        self.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.session_file = self.training_data_dir / f"session_{self.session_id}.json"
        
        self.session_data = {
            'session_id': self.session_id,
            'pairs_processed': 0,
            'corrections_found': 0,
            'rules_generated': 0,
            'approved_dishes_stored': 0,
            'pairs': [],
            'all_corrections': [],
            'generated_rules': []
        }
    
    def ingest_document_pair(
        self, 
        original_path: str, 
        redlined_path: str
    ) -> Dict:
        """
        Ingest and analyze a single document pair.
        """
        print(f"\nProcessing pair:")
        print(f"  Original: {os.path.basename(original_path)}")
        print(f"  Redlined: {os.path.basename(redlined_path)}")
        
        analysis = self.analyzer.load_document_pair(original_path, redlined_path)
        
        # Update session stats
        self.session_data['pairs_processed'] += 1
        self.session_data['corrections_found'] += len(analysis['text_corrections'])
        self.session_data['pairs'].append({
            'original': original_path,
            'redlined': redlined_path,
            'corrections': len(analysis['text_corrections'])
        })
        
        # Store all corrections
        self.session_data['all_corrections'].extend(analysis['text_corrections'])
        
        # Store corrections in dish database (both allergens AND full descriptions)
        if DISH_DB_AVAILABLE:
            restaurant = extract_restaurant(original_path)
            dishes_stored = 0
            
            for correction in analysis['text_corrections']:
                category = correction.get('category', '')
                
                # Skip swapped items - these are not real corrections
                if category == 'swapped_item':
                    continue
                
                original_line = correction.get('original', '')
                corrected_line = correction.get('corrected', '')
                
                # Store allergen corrections
                if category == 'allergen':
                    for word_diff in correction.get('word_diffs', []):
                        orig_codes = word_diff.get('original_words', '')
                        corr_codes = word_diff.get('corrected_words', '')
                        
                        result = store_allergen_correction(
                            dish_line=original_line,
                            original_codes=orig_codes,
                            corrected_codes=corr_codes,
                            restaurant=restaurant
                        )
                        if result:
                            dishes_stored += 1
                
                # Store COMPLETE dish descriptions for spelling/general corrections
                # This captures ingredient changes like "pork chorizo" → "bacon"
                elif category in ('spelling', 'general', 'diacritics'):
                    # Only store if it looks like a menu item (has commas = has ingredients)
                    if ', ' in corrected_line and len(corrected_line) > 20:
                        result = learn_dish_from_correction(
                            original_line=original_line,
                            corrected_line=corrected_line,
                            restaurant=restaurant
                        )
                        if result:
                            dishes_stored += 1
                
                # Store terminology corrections with dish context
                # So next time we see "Red Paloma", we know "crust" → "rim"
                elif category == 'terminology':
                    context = correction.get('context', {})
                    dish_name = context.get('dish_name')
                    
                    if dish_name:
                        for word_diff in correction.get('word_diffs', []):
                            orig_term = word_diff.get('original_words', '')
                            corr_term = word_diff.get('corrected_words', '')
                            
                            if orig_term and corr_term:
                                result = store_dish_terminology_correction(
                                    dish_name=dish_name,
                                    original_term=orig_term,
                                    corrected_term=corr_term,
                                    restaurant=restaurant,
                                    context_paragraph=context.get('paragraph')
                                )
                                if result:
                                    dishes_stored += 1
                                    print(f"    Stored terminology '{orig_term}' → '{corr_term}' for dish: {dish_name}")
            
            if dishes_stored > 0:
                print(f"  Stored {dishes_stored} corrections to dish database")
            
            # NEW: Store ALL dishes from the approved (redlined) document
            # This builds the master catalog of approved dishes
            approved_dishes_stored = self._store_all_approved_dishes(
                redlined_path, 
                restaurant
            )
            if approved_dishes_stored > 0:
                print(f"  Stored {approved_dishes_stored} approved dishes to master catalog")
                self.session_data['approved_dishes_stored'] += approved_dishes_stored
        
        print(f"  Found {len(analysis['text_corrections'])} text corrections")
        print(f"  Found {len(analysis['formatting_corrections'])} formatting corrections")
        
        return analysis
    
    def _store_all_approved_dishes(
        self, 
        redlined_path: str, 
        restaurant: str
    ) -> int:
        """
        Extract and store ALL dishes from an approved (redlined) document.
        
        This builds the master catalog of approved dishes with their:
        - Names
        - Full descriptions (ingredients)
        - Allergen codes
        - Prices
        - Restaurant
        - Menu date
        - Menu type
        
        Args:
            redlined_path: Path to the approved document
            restaurant: Restaurant identifier
            
        Returns:
            Number of dishes stored
        """
        if not DISH_DB_AVAILABLE:
            return 0
        
        doc = Document(redlined_path)
        
        # Extract menu date and type from filename
        menu_date = extract_menu_date(redlined_path)
        menu_type = extract_menu_type(redlined_path)
        
        dishes_stored = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip empty lines
            if not text or len(text) < 5:
                continue
            
            # Skip lines that don't look like menu items
            # Menu items typically have: comma-separated ingredients, or are short dish names
            
            # Skip headers (usually ALL CAPS or very short)
            if text.isupper() and len(text.split()) <= 3:
                continue
            
            # Skip warning text
            if 'consuming raw' in text.lower() or 'foodborne' in text.lower():
                continue
            
            # Skip page numbers, template headers, etc.
            if text.lower().startswith(('page', 'menu', 'restaurant', 'venue')):
                continue
            
            # Only process lines that look like menu items
            # They typically have commas (ingredients) or a price at the end
            has_comma = ', ' in text
            has_price = bool(re.search(r'\s+\d+(?:\|\d+)?\s*$', text))
            has_allergens = bool(re.search(r'\s+[A-Z,]+\s*$', text))
            
            # Must have at least one indicator of being a menu item
            if not (has_comma or has_price or has_allergens):
                continue
            
            # Store this dish
            result = store_approved_dish(
                dish_line=text,
                restaurant=restaurant,
                menu_date=menu_date,
                menu_type=menu_type,
                source_file=os.path.basename(redlined_path)
            )
            
            if result:
                dishes_stored += 1
        
        return dishes_stored
    
    def ingest_directory_pairs(
        self, 
        directory: str,
        original_pattern: str = "*original*.docx",
        redlined_pattern: str = "*redlined*.docx"
    ):
        """
        Ingest all document pairs from a directory.
        
        Expects files to follow naming convention:
        - original files contain 'original' in name
        - redlined files contain 'redlined' in name and match original name
        """
        dir_path = Path(directory)
        
        original_files = sorted(dir_path.glob(original_pattern))
        
        print(f"\nScanning directory: {directory}")
        print(f"Found {len(original_files)} original files")
        
        for orig_file in original_files:
            # Try to find matching redlined file
            base_name = orig_file.stem.replace('original', '').replace('_', '').strip()
            
            # Look for redlined version
            redlined_candidates = list(dir_path.glob(f"*{base_name}*redlined*.docx"))
            
            if not redlined_candidates:
                print(f"\nWarning: No redlined version found for {orig_file.name}")
                continue
            
            redlined_file = redlined_candidates[0]
            
            # Process the pair
            self.ingest_document_pair(str(orig_file), str(redlined_file))
    
    def generate_rules(self, min_occurrences: int = 2, existing_rules_path: str = None) -> List[Dict]:
        """
        Generate rules from all accumulated corrections.
        
        Args:
            min_occurrences: Minimum times a pattern must appear to become a rule
            existing_rules_path: Path to existing rules file to avoid duplicates
        """
        print(f"\n{'='*60}")
        print("GENERATING RULES FROM TRAINING DATA")
        print(f"{'='*60}")
        
        # Filter out swapped items from corrections before counting
        valid_corrections = [
            c for c in self.session_data['all_corrections']
            if c.get('category') != 'swapped_item'
        ]
        swapped_count = len(self.session_data['all_corrections']) - len(valid_corrections)
        
        print(f"Total corrections analyzed: {len(valid_corrections)}")
        if swapped_count > 0:
            print(f"Filtered out {swapped_count} swapped menu items (not real corrections)")
        
        # Load existing rules to avoid duplicates
        existing_rules = []
        if existing_rules_path and os.path.exists(existing_rules_path):
            try:
                with open(existing_rules_path, 'r') as f:
                    data = json.load(f)
                    existing_rules = data.get('rules', [])
                    existing_learned = [r for r in existing_rules if r.get('rule_id', '').startswith('LEARNED')]
                    if existing_learned:
                        print(f"Found {len(existing_learned)} existing learned rules (will skip duplicates)")
            except Exception as e:
                print(f"Note: Could not load existing rules: {e}")
        
        rules = self.rule_generator.generate_rules_from_corrections(
            valid_corrections,
            min_occurrences=min_occurrences,
            existing_rules=existing_rules
        )
        
        self.session_data['generated_rules'] = rules
        self.session_data['rules_generated'] = len(rules)
        
        print(f"Generated {len(rules)} NEW rules (not counting duplicates)")
        
        return rules
    
    def save_rules_to_file(self, output_path: str = None):
        """
        Save generated rules to a JSON file.
        """
        if output_path is None:
            output_path = str(self.training_data_dir / f"learned_rules_{self.session_id}.json")
        
        rules_data = {
            'generated_at': datetime.now().isoformat(),
            'session_id': self.session_id,
            'pairs_processed': self.session_data['pairs_processed'],
            'rules': self.session_data['generated_rules']
        }
        
        with open(output_path, 'w') as f:
            json.dump(rules_data, f, indent=2)
        
        print(f"\nRules saved to: {output_path}")
        return output_path
    
    def merge_with_existing_rules(self, existing_rules_path: str, output_path: str):
        """
        Merge generated rules with existing SOP rules.
        """
        # Load existing rules
        with open(existing_rules_path, 'r') as f:
            existing_data = json.load(f)
        
        # Add new rules
        existing_data['rules'].extend(self.session_data['generated_rules'])
        
        # Save merged rules
        with open(output_path, 'w') as f:
            json.dump(existing_data, f, indent=2)
        
        print(f"\nMerged rules saved to: {output_path}")
    
    def generate_optimized_prompt(self, current_prompt: str) -> str:
        """
        Generate an optimized prompt based on training data.
        """
        optimizer = PromptOptimizer(current_prompt)
        optimizer.add_training_examples(self.session_data['all_corrections'])
        
        enhanced_prompt = optimizer.generate_enhanced_prompt()
        
        # Save to file
        prompt_file = self.training_data_dir / f"optimized_prompt_{self.session_id}.txt"
        with open(prompt_file, 'w') as f:
            f.write(enhanced_prompt)
        
        print(f"\nOptimized prompt saved to: {prompt_file}")
        
        return enhanced_prompt
    
    def save_session(self):
        """
        Save the current training session data.
        """
        with open(self.session_file, 'w') as f:
            json.dump(self.session_data, f, indent=2)
        
        print(f"\nSession data saved to: {self.session_file}")
    
    def print_summary(self):
        """
        Print a summary of the training session.
        """
        print(f"\n{'='*60}")
        print("TRAINING SESSION SUMMARY")
        print(f"{'='*60}")
        print(f"Session ID: {self.session_id}")
        print(f"Document pairs processed: {self.session_data['pairs_processed']}")
        print(f"Total corrections found: {self.session_data['corrections_found']}")
        print(f"Rules generated: {self.session_data['rules_generated']}")
        print(f"Approved dishes added to catalog: {self.session_data['approved_dishes_stored']}")
        
        # Category breakdown
        if self.session_data['all_corrections']:
            categories = Counter(c.get('category', 'unknown') 
                                for c in self.session_data['all_corrections'])
            
            print("\nCorrections by category:")
            for category, count in categories.most_common():
                print(f"  {category}: {count}")
        
        # Dish allergen database stats
        if DISH_DB_AVAILABLE:
            try:
                dish_stats = get_dish_db_stats()
                if dish_stats.get('total_dishes', 0) > 0:
                    print("\nDish Allergen Database:")
                    print(f"  Total dishes: {dish_stats['total_dishes']}")
                    if dish_stats.get('by_restaurant'):
                        print("  By restaurant:")
                        for restaurant, count in dish_stats['by_restaurant'].items():
                            print(f"    {restaurant}: {count}")
            except Exception as e:
                print(f"\nNote: Could not load dish database stats: {e}")
        
        print(f"\n{'='*60}\n")


# CLI Interface
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Train the menu redliner from document pairs"
    )
    
    parser.add_argument(
        '--directory',
        '-d',
        help='Directory containing document pairs',
        required=True
    )
    
    parser.add_argument(
        '--min-occurrences',
        '-m',
        type=int,
        default=2,
        help='Minimum occurrences for a pattern to become a rule (default: 2)'
    )
    
    parser.add_argument(
        '--merge-rules',
        help='Path to existing rules file to merge with',
        default=None
    )
    
    parser.add_argument(
        '--optimize-prompt',
        action='store_true',
        help='Generate optimized AI prompt'
    )
    
    args = parser.parse_args()
    
    # Initialize pipeline
    pipeline = TrainingPipeline()
    
    # Ingest document pairs
    pipeline.ingest_directory_pairs(args.directory)
    
    # Generate rules (pass existing rules path to avoid duplicates)
    pipeline.generate_rules(
        min_occurrences=args.min_occurrences,
        existing_rules_path=args.merge_rules
    )
    
    # Save rules
    rules_file = pipeline.save_rules_to_file()
    
    # Merge with existing rules if requested
    if args.merge_rules:
        output_path = args.merge_rules.replace('.json', '_updated.json')
        pipeline.merge_with_existing_rules(args.merge_rules, output_path)
    
    # Generate optimized prompt if requested
    if args.optimize_prompt:
        from ai_corrector import AICorrector
        corrector = AICorrector()
        pipeline.generate_optimized_prompt(corrector.system_prompt)
    
    # Save session and print summary
    pipeline.save_session()
    pipeline.print_summary()
    
    print("\n✓ Training complete!")
    print(f"  Review generated rules in: {rules_file}")
