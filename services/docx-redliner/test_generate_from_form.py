#!/usr/bin/env python3
"""Tests for generated form DOCX menu formatting."""

import os
import tempfile

from docx import Document

from generate_from_form import populate_template


def test_populate_template_preserves_strong_dish_name_runs():
    template = Document()
    template.add_paragraph("MENU")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as template_file:
        template.save(template_file.name)
        template_path = template_file.name

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_file:
        output_path = output_file.name

    try:
        populate_template(
            template_path,
            {
                "projectName": "Dinner Menu",
                "property": "Test Property",
                "size": "8.5 x 11 inches",
                "orientation": "Portrait",
                "dateNeeded": "2026-06-04",
                "menuContent": "Guacamole - fresh avocado / lime / cilantro D,G 12",
                "menuContentHtml": "<p><strong>Guacamole</strong> - fresh avocado / lime / cilantro D,G 12</p>",
                "allergens": "",
                "footerText": "",
                "shouldAddRawNotice": False,
            },
            output_path,
        )

        generated = Document(output_path)
        menu_paragraph = next(
            paragraph
            for paragraph in generated.paragraphs
            if paragraph.text.startswith("Guacamole")
        )

        assert menu_paragraph.runs[0].text == "Guacamole"
        assert menu_paragraph.runs[0].bold is True
        assert "".join(run.text for run in menu_paragraph.runs[1:]).startswith(
            " - fresh avocado"
        )
        assert all(run.bold is not True for run in menu_paragraph.runs[1:])
    finally:
        for path in (template_path, output_path):
            if os.path.exists(path):
                os.unlink(path)
