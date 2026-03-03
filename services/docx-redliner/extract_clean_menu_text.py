#!/usr/bin/env python3
"""
Extract menu text from DOCX and produce a "cleaned" version that attempts
to remove redline/deletion artifacts.

Usage:
    python extract_clean_menu_text.py <docx_path> [--mode clean|unapproved]
"""

import json
import sys
from html import escape
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

BOUNDARY_MARKERS = [
    "Please drop the menu content below on page 2",
    "MENU",
]


# ── Module-level helpers ─────────────────────────────────────────────────────

def local_name(tag):
    """Extract the local name from an XML tag, stripping namespace."""
    if not isinstance(tag, str):
        return ""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def run_is_in_deleted_change(run):
    """Return True if the run sits inside a tracked deletion (w:del / w:moveFrom)."""
    node = run._r
    while node is not None:
        if local_name(node.tag) in {"del", "moveFrom"}:
            return True
        node = node.getparent()
    return False


def run_is_in_inserted_change(run):
    """Return True if the run sits inside a tracked insertion (w:ins / w:moveTo)."""
    node = run._r
    while node is not None:
        if local_name(node.tag) in {"ins", "moveTo"}:
            return True
        node = node.getparent()
    return False


def run_has_highlight(run):
    """Return True if the run has a highlight color applied."""
    try:
        return run.font.highlight_color is not None and run.font.highlight_color != WD_COLOR_INDEX.AUTO
    except Exception:
        return False


# ── Boundary detection ────────────────────────────────────────────────────────

def find_boundary_index(paragraphs):
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "MENU":
            return i
        if BOUNDARY_MARKERS[0] in text:
            return i
    return None


# ── Clean mode (approved DOCX – strips redlines) ─────────────────────────────

def paragraph_clean_text(paragraph):
    """
    Best-effort cleaner:
    - Removes explicitly struck-through text (common manual redline deletion style).
    - Keeps inserted/highlighted text.
    - Removes Word tracked deletions (w:del / w:moveFrom) and keeps
      tracked insertions (w:ins / w:moveTo).
    """
    if not paragraph.runs:
        return paragraph.text

    cleaned_parts = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        if run_is_in_deleted_change(run):
            continue
        if run.font and run.font.strike:
            continue
        cleaned_parts.append(text)

    return "".join(cleaned_parts)


def paragraph_clean_html(paragraph):
    """
    Build cleaned HTML paragraph preserving basic inline formatting:
    bold, italic, underline. Tracked/manual deletions are removed.
    """
    fragments = []
    for run in paragraph.runs:
        if run_is_in_deleted_change(run):
            continue
        if run.font and run.font.strike:
            continue

        text = run.text or ""
        if not text:
            continue

        chunk = escape(text).replace("\n", "<br>")
        if run.bold:
            chunk = f"<strong>{chunk}</strong>"
        if run.italic:
            chunk = f"<em>{chunk}</em>"
        if run.underline:
            chunk = f"<u>{chunk}</u>"
        fragments.append(chunk)

    if not fragments:
        return "<p><br></p>"
    return f"<p>{''.join(fragments)}</p>"


# ── Unapproved mode (preserves existing redlines) ────────────────────────────

def paragraph_unapproved_text(paragraph):
    """
    Return ALL visible text including deletions (struck-through / tracked deletes).
    This is the full text the chef sees when reviewing the unapproved document.
    """
    if not paragraph.runs:
        return paragraph.text
    parts = []
    for run in paragraph.runs:
        text = run.text or ""
        if text:
            parts.append(text)
    return "".join(parts)


def paragraph_unapproved_html(paragraph):
    """
    Build HTML that preserves existing redlines:
    - Tracked deletions / strikethrough → <span class="existing-del">
    - Tracked insertions / highlighted → <span class="existing-ins">
    - Everything else rendered normally with bold/italic/underline.
    """
    fragments = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue

        is_deleted = run_is_in_deleted_change(run) or (run.font and run.font.strike)
        is_inserted = run_is_in_inserted_change(run) or run_has_highlight(run)

        chunk = escape(text).replace("\n", "<br>")

        # Apply inline formatting
        if run.bold:
            chunk = f"<strong>{chunk}</strong>"
        if run.italic:
            chunk = f"<em>{chunk}</em>"
        if run.underline and not is_deleted:
            chunk = f"<u>{chunk}</u>"

        # Wrap with redline class (deletion takes priority if both flags set)
        if is_deleted:
            chunk = f'<span class="existing-del">{chunk}</span>'
        elif is_inserted:
            chunk = f'<span class="existing-ins">{chunk}</span>'

        fragments.append(chunk)

    if not fragments:
        return "<p><br></p>"
    return f"<p>{''.join(fragments)}</p>"


def paragraph_annotation_ranges(paragraph):
    """
    Return character-level annotation ranges for a paragraph.
    Each entry: { start, end, type } where type is 'del' or 'ins'.
    Offsets are relative to the full visible text (including deletions).
    """
    annotations = []
    offset = 0
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue

        length = len(text)
        is_deleted = run_is_in_deleted_change(run) or (run.font and run.font.strike)
        is_inserted = run_is_in_inserted_change(run) or run_has_highlight(run)

        if is_deleted:
            annotations.append({"start": offset, "end": offset + length, "type": "del"})
        elif is_inserted:
            annotations.append({"start": offset, "end": offset + length, "type": "ins"})

        offset += length

    return annotations


# ── Main extraction ───────────────────────────────────────────────────────────

def extract_texts(docx_path, mode="clean"):
    doc = Document(docx_path)
    boundary_index = find_boundary_index(doc.paragraphs)

    if boundary_index is None:
        source_paragraphs = doc.paragraphs
    else:
        source_paragraphs = doc.paragraphs[boundary_index + 1 :]

    raw_lines = []
    cleaned_lines = []
    cleaned_html_paragraphs = []

    # Unapproved mode extras
    visible_lines = []
    unapproved_html_paragraphs = []
    all_annotations = []

    for paragraph in source_paragraphs:
        raw_text = paragraph.text

        # Skip instruction/template markers that may still appear.
        if raw_text.strip() == "MENU" or "Please drop the menu content below" in raw_text:
            continue

        raw_lines.append(raw_text)

        if mode == "unapproved":
            visible_lines.append(paragraph_unapproved_text(paragraph))
            unapproved_html_paragraphs.append(paragraph_unapproved_html(paragraph))
            all_annotations.append(paragraph_annotation_ranges(paragraph))
        else:
            cleaned_text = paragraph_clean_text(paragraph)
            cleaned_lines.append(cleaned_text)
            cleaned_html_paragraphs.append(paragraph_clean_html(paragraph))

    # Trim trailing blank lines
    while raw_lines and not raw_lines[-1].strip():
        raw_lines.pop()

    if mode == "unapproved":
        while visible_lines and not visible_lines[-1].strip():
            visible_lines.pop()
            if all_annotations:
                all_annotations.pop()
            if unapproved_html_paragraphs:
                unapproved_html_paragraphs.pop()

        return {
            "menu_content": "\n".join(raw_lines),
            "visible_text": "\n".join(visible_lines),
            "unapproved_html": "".join(unapproved_html_paragraphs),
            "annotations": all_annotations,
        }

    while cleaned_lines and not cleaned_lines[-1].strip():
        cleaned_lines.pop()

    return {
        "menu_content": "\n".join(raw_lines),
        "cleaned_menu_content": "\n".join(cleaned_lines),
        "cleaned_menu_html": "".join(cleaned_html_paragraphs),
    }


def main():
    args = sys.argv[1:]
    mode = "clean"

    # Parse --mode flag
    if "--mode" in args:
        idx = args.index("--mode")
        if idx + 1 < len(args):
            mode = args[idx + 1]
            args = args[:idx] + args[idx + 2:]

    if len(args) != 1:
        print(json.dumps({"error": "Usage: python extract_clean_menu_text.py <docx_path> [--mode clean|unapproved]"}))
        sys.exit(1)

    docx_path = args[0]
    try:
        payload = extract_texts(docx_path, mode=mode)
        print(json.dumps(payload))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
