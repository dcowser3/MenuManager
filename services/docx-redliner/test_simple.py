from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Create a simple test document
doc = Document()
para = doc.add_paragraph()

# Add text with red strikethrough (deletion)
run1 = para.add_run("This is deleted")
run1.font.strike = True
run1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# Add normal space
para.add_run(" ")

# Add text with yellow highlight (addition)
run2 = para.add_run("This is added")
run2.font.highlight_color = WD_COLOR_INDEX.YELLOW

doc.save('/tmp/formatting_test.docx')
print("Test document created: /tmp/formatting_test.docx")
print("Please open it in Word to verify formatting appears correctly")
