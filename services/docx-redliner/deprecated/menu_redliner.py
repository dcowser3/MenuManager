"""
Menu Document Redliner
======================
This module processes Word documents containing menu content and applies
AI-generated corrections with tracked changes (strikethrough for deletions,
highlight for additions) while preserving all original formatting.

The document is expected to have a boundary marker that separates the
template content from the menu content. Only content after the boundary
marker is processed.

Special handling for Prix Fixe menus:
- Detects prix fixe/tasting menus by keywords
- Automatically adds course numbers (1, 2, 3, 4...) before section headers
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import diff_match_patch as dmp_module
from typing import List, Tuple, Optional
import re
import difflib
import os
from pathlib import Path

# The text that separates the template from the menu
BOUNDARY_MARKER = "Please drop the menu content below on page 2."

# Keywords that indicate a prix fixe / tasting menu
PRIX_FIXE_KEYWORDS = [
    "prix fixe", "pre-fix", "prefix", "prix-fixe",
    "tasting menu", "tasting experience",
    "course menu", "multi-course", "multicourse",
    "degustation", "chef's menu", "chef's table"
]

# Patterns for course section headers (creative names with subtitles)
# Examples: "The Spark – "El Primer Encuentro"", "The Connection – "El Abrazo""
# Note: Use explicit Unicode escapes to ensure proper matching
COURSE_HEADER_PATTERNS = [
    # Pattern: "The X – "Spanish Subtitle"" (using en-dash U+2013 and curly quotes U+201C/U+201D)
    r'^The\s+\w+\s*[\u2013\u2014-]\s*[\u201c\u201d"][^\u201c\u201d"]+[\u201c\u201d"]',
    # Pattern: "Course Name – Subtitle" with curly or straight quotes
    r'^[A-Z][a-z]+(?:\s+[A-Z]?[a-z]+)*\s*[\u2013\u2014-]\s*[\u201c\u201d"][^\u201c\u201d"]+[\u201c\u201d"]',
    # Pattern: "COURSE X" or "Course X" followed by name
    r'^(?:COURSE|Course)\s+(?:One|Two|Three|Four|Five|Six|Seven|Eight|\d+)',
]

# Default allergen codes (from RSH SOP) - used when document has no legend
DEFAULT_ALLERGEN_CODES = {
    'D': 'Dairy',
    'G': 'Gluten',
    'N': 'Nuts',
    'S': 'Shellfish',
    'V': 'Vegetarian',
    'VG': 'Vegan',
}


def detect_allergen_legend(paragraphs: list) -> dict:
    """
    Detect and parse an allergen legend from document paragraphs.

    Looks for patterns like:
    - "C crustaceans | CE celery | D dairy | E egg | F fish..."
    - "D=Dairy | G=Gluten | N=Nuts..."
    - Lines containing multiple "CODE word" pairs separated by |

    Args:
        paragraphs: List of paragraph objects from the document

    Returns:
        Dict mapping allergen codes to their meanings, or empty dict if no legend found
    """
    # Check the last 15 paragraphs (legends usually at bottom)
    check_paragraphs = paragraphs[-15:] if len(paragraphs) > 15 else paragraphs

    for para in check_paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Look for patterns with | separators containing allergen definitions
        # Pattern: multiple entries separated by |, each with CODE + word
        if '|' in text:
            # Count potential allergen entries (CODE followed by word)
            # Each entry looks like "C crustaceans" or "CE celery" or "D=dairy"
            entries = text.split('|')

            if len(entries) >= 5:  # Legend should have many allergen types
                allergen_codes = {}
                valid_entries = 0

                for entry in entries:
                    entry = entry.strip()
                    # Match patterns like "C crustaceans", "CE celery", "D=dairy", "VG vegan"
                    match = re.match(r'^([A-Z]{1,3})\s*[=\-:]?\s*(\w+(?:\s+\w+)?)', entry, re.IGNORECASE)
                    if match:
                        code = match.group(1).upper()
                        name = match.group(2).lower()
                        # Validate it looks like an allergen (not random text)
                        allergen_keywords = [
                            'dairy', 'gluten', 'nuts', 'nut', 'shellfish', 'fish',
                            'vegetarian', 'vegan', 'egg', 'celery', 'crustacean',
                            'sesame', 'soya', 'soy', 'mustard', 'lupin', 'mollusc',
                            'sulphite', 'sulfite', 'peanut', 'tree'
                        ]
                        if any(keyword in name for keyword in allergen_keywords) or len(code) <= 2:
                            allergen_codes[code] = name.title()
                            valid_entries += 1

                # If we found enough valid allergen entries, this is likely the legend
                if valid_entries >= 5:
                    print(f"  Detected document allergen legend with {len(allergen_codes)} codes: {', '.join(sorted(allergen_codes.keys()))}")
                    return allergen_codes

    return {}


def get_allergen_codes_for_document(paragraphs: list) -> dict:
    """
    Get the allergen codes to use for a document.

    First tries to detect a document-specific allergen legend.
    Falls back to DEFAULT_ALLERGEN_CODES if none found.

    Args:
        paragraphs: List of paragraph objects from the document

    Returns:
        Dict mapping allergen codes to their meanings
    """
    # Try to detect document-specific legend
    doc_legend = detect_allergen_legend(paragraphs)

    if doc_legend:
        return doc_legend

    # Fall back to default SOP codes
    return DEFAULT_ALLERGEN_CODES.copy()


class MenuRedliner:
    """
    Handles the processing of menu documents with AI corrections
    and formatted diff tracking.
    """
    
    def __init__(self, boundary_marker: str = BOUNDARY_MARKER):
        """
        Initialize the MenuRedliner.
        
        Args:
            boundary_marker: The text that marks the start of menu content
        """
        self.boundary_marker = boundary_marker
        self.dmp = dmp_module.diff_match_patch()
    
    def is_prix_fixe_menu(self, paragraphs: List) -> bool:
        """
        Detect if the menu is a prix fixe / tasting menu by scanning
        for keywords in the first several paragraphs.
        
        Args:
            paragraphs: List of paragraph objects to scan
            
        Returns:
            True if this appears to be a prix fixe menu
        """
        # Check the first 10 paragraphs for prix fixe keywords
        text_to_check = " ".join(p.text.lower() for p in paragraphs[:10])
        
        for keyword in PRIX_FIXE_KEYWORDS:
            if keyword in text_to_check:
                print(f"  Detected prix fixe menu (keyword: '{keyword}')")
                return True
        
        return False
    
    def is_course_header(self, text: str) -> bool:
        """
        Check if a paragraph text matches a course section header pattern.
        
        Args:
            text: The paragraph text to check
            
        Returns:
            True if this looks like a course section header
        """
        text = text.strip()
        if not text:
            return False
        
        for pattern in COURSE_HEADER_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def add_course_numbers(self, doc, paragraphs_to_process: List) -> int:
        """
        Add course numbers before section headers in a prix fixe menu.
        Inserts new paragraphs with just the course number.
        
        Args:
            doc: The Document object
            paragraphs_to_process: List of menu paragraphs
            
        Returns:
            Number of course numbers added
        """
        course_number = 0
        paragraphs_with_numbers = []
        
        # First pass: identify which paragraphs are course headers
        for para in paragraphs_to_process:
            if self.is_course_header(para.text):
                course_number += 1
                paragraphs_with_numbers.append((para, course_number))
        
        if course_number == 0:
            print("  No course headers found to number")
            return 0
        
        print(f"  Found {course_number} course sections to number")
        
        # Second pass: insert course numbers
        # We need to insert from bottom to top to preserve indices
        for para, num in reversed(paragraphs_with_numbers):
            # Find the paragraph's position in the document body
            body = doc._body._body
            para_element = para._element
            
            # Create a new paragraph for the course number
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Create new paragraph element
            new_p = OxmlElement('w:p')
            
            # Create paragraph properties for centering (matching the style)
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            new_p.append(pPr)
            
            # Create run with the number
            run = OxmlElement('w:r')
            
            # Run properties (bold, same font)
            rPr = OxmlElement('w:rPr')
            bold = OxmlElement('w:b')
            rPr.append(bold)
            run.append(rPr)
            
            # Text content
            text_elem = OxmlElement('w:t')
            text_elem.text = str(num)
            run.append(text_elem)
            
            new_p.append(run)
            
            # Insert the new paragraph before the course header
            para_element.addprevious(new_p)
            
            # Also add yellow highlight to mark it as an addition
            from docx.oxml.ns import nsmap
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'yellow')
            rPr.append(highlight)
        
        return course_number
    
    def find_run_at_index(self, original_runs: List, original_text_len: int, target_index: int):
        """
        Helper function to find which original run (and its style)
        corresponds to a character index in the paragraph's plain text.

        Args:
            original_runs: List of runs from the original paragraph
            original_text_len: Total length of the original text
            target_index: Character index to find

        Returns:
            The run that contains the character at target_index
        """
        cumulative_len = 0
        last_text_run = None  # Track the last run that had actual text

        for run in original_runs:
            run_len = len(run.text)
            if run_len > 0:
                last_text_run = run
            if cumulative_len + run_len > target_index:
                return run
            cumulative_len += run_len

        # Fallback: prefer the first run with content (body text style)
        # over the last run (which might be bold pricing)
        # This prevents inheriting bold from prices at end of lines
        for run in original_runs:
            if run.text.strip():
                return run

        return original_runs[0] if original_runs else None
    
    def build_char_format_map(self, original_runs: List) -> List[dict]:
        """
        Build a per-character formatting map from original runs.
        Each entry contains the formatting for that character position.
        """
        format_map = []
        for run in original_runs:
            run_format = {
                'name': run.font.name,
                'size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None,
            }
            for _ in run.text:
                format_map.append(run_format)
        return format_map

    def apply_formatted_diffs(self, para, diffs: List[Tuple[int, str]]):
        """
        Rebuilds a paragraph run-by-run based on diffs,
        preserving original styles and applying diff formatting.

        Args:
            para: The paragraph object to modify
            diffs: List of (operation, text) tuples from diff_match_patch
        """
        # 1. Save original runs and build a per-character format map
        original_runs = list(para.runs)
        original_full_text = "".join(r.text for r in original_runs)
        char_format_map = self.build_char_format_map(original_runs)

        if not original_runs:
            # If there are no runs, just create new ones
            for op, text in diffs:
                if not text:
                    continue
                new_run = para.add_run(text)
                if op == dmp_module.diff_match_patch.DIFF_DELETE:
                    new_run.font.strike = True
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif op == dmp_module.diff_match_patch.DIFF_INSERT:
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return

        # 2. Clear the paragraph's existing content
        # We need to remove all runs while keeping the paragraph
        for run in para.runs:
            run._element.getparent().remove(run._element)
        
        current_text_index = 0

        # Get default format from first run (for inserts)
        default_format = char_format_map[0] if char_format_map else {}

        # 3. Iterate through diffs
        for op, text in diffs:
            if not text:
                continue

            if op == dmp_module.diff_match_patch.DIFF_EQUAL:
                # For EQUAL text, preserve exact per-character formatting
                # Split into runs based on formatting changes
                i = 0
                while i < len(text):
                    char_idx = current_text_index + i
                    if char_idx < len(char_format_map):
                        current_format = char_format_map[char_idx]
                    else:
                        current_format = default_format

                    # Find consecutive characters with same format
                    run_text = text[i]
                    j = i + 1
                    while j < len(text):
                        next_idx = current_text_index + j
                        if next_idx < len(char_format_map):
                            next_format = char_format_map[next_idx]
                        else:
                            next_format = default_format

                        # Check if format changed (specifically bold)
                        if next_format.get('bold') != current_format.get('bold'):
                            break
                        run_text += text[j]
                        j += 1

                    # Create run with correct formatting
                    new_run = para.add_run(run_text)
                    if current_format.get('name'):
                        new_run.font.name = current_format['name']
                    if current_format.get('size'):
                        new_run.font.size = current_format['size']
                    if current_format.get('bold') is not None:
                        new_run.font.bold = current_format['bold']
                    if current_format.get('italic') is not None:
                        new_run.font.italic = current_format['italic']
                    if current_format.get('underline') is not None:
                        new_run.font.underline = current_format['underline']
                    if current_format.get('color'):
                        new_run.font.color.rgb = current_format['color']

                    i = j

                current_text_index += len(text)

            elif op == dmp_module.diff_match_patch.DIFF_DELETE:
                # For DELETE, apply red strikethrough
                new_run = para.add_run(text)
                # Copy font name/size from original position
                if current_text_index < len(char_format_map):
                    fmt = char_format_map[current_text_index]
                    if fmt.get('name'):
                        new_run.font.name = fmt['name']
                    if fmt.get('size'):
                        new_run.font.size = fmt['size']
                    if fmt.get('bold') is not None:
                        new_run.font.bold = fmt['bold']
                new_run.font.strike = True
                new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                current_text_index += len(text)

            elif op == dmp_module.diff_match_patch.DIFF_INSERT:
                # For INSERT, apply yellow highlight, NO bold
                new_run = para.add_run(text)
                # Only copy font name/size, NOT bold
                if default_format.get('name'):
                    new_run.font.name = default_format['name']
                if default_format.get('size'):
                    new_run.font.size = default_format['size']
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # Don't advance current_text_index for inserts
    
    def paragraph_has_existing_redlines(self, para) -> bool:
        """
        Check if a paragraph has ANY existing redlines (strikethrough or highlight).
        """
        for run in para.runs:
            if run.font.strike:
                return True
            if run.font.highlight_color:
                return True
        return False

    def paragraph_has_mixed_bold(self, para) -> bool:
        """
        Check if a paragraph has mixed bold formatting (some runs bold, some not).
        This typically means dish name is bold but description is not.
        We should skip these to avoid corrupting the formatting.
        """
        has_bold = False
        has_non_bold = False

        for run in para.runs:
            # Skip empty runs
            if not run.text.strip():
                continue
            if run.font.bold:
                has_bold = True
            else:
                has_non_bold = True

            # If we found both, it's mixed
            if has_bold and has_non_bold:
                return True

        return False

    def process_paragraph(self, para, ai_correction_func) -> bool:
        """
        Process a single paragraph with AI correction and diff application.

        IMPORTANT: We skip paragraphs in these cases:
        1. Already has redlines (strikethrough/highlight) - would destroy existing edits
        2. Has mixed bold formatting (dish name bold, description not) - would corrupt styling

        Args:
            para: The paragraph object to process
            ai_correction_func: Function that takes original text and returns corrected text

        Returns:
            True if the paragraph was modified, False otherwise
        """
        original_text = para.text
        if not original_text.strip():
            # Skip empty paragraphs
            return False

        # CRITICAL: Skip paragraphs that already have ANY redlines
        # Rebuilding runs would destroy the existing redline formatting
        if self.paragraph_has_existing_redlines(para):
            return False

        # Get AI correction
        corrected_text = ai_correction_func(original_text)

        # If no changes, skip processing
        if original_text == corrected_text:
            return False

        # Compute the difference using WORD-LEVEL diffs
        diffs = self._word_level_diffs(original_text, corrected_text)

        # Apply the formatted diffs
        self.apply_formatted_diffs(para, diffs)

        return True

    def _tokenize(self, text: str) -> List[str]:
        """Split text into tokens: words, whitespace, and punctuation.
        This ensures diffs happen on word boundaries.
        """
        return re.findall(r"\w+|\s+|[^\w\s]", text)

    def _word_level_diffs(self, original_text: str, corrected_text: str) -> List[Tuple[int, str]]:
        """
        Produce diffs at word/token level using difflib, then map to
        diff_match_patch-style tuples. This forces whole-word replacements.

        IMPORTANT: Whitespace-only changes are ignored to avoid cluttering redlines.
        """
        DIFF_EQUAL = dmp_module.diff_match_patch.DIFF_EQUAL
        DIFF_DELETE = dmp_module.diff_match_patch.DIFF_DELETE
        DIFF_INSERT = dmp_module.diff_match_patch.DIFF_INSERT

        orig_tokens = self._tokenize(original_text)
        corr_tokens = self._tokenize(corrected_text)

        sm = difflib.SequenceMatcher(a=orig_tokens, b=corr_tokens)
        diffs: List[Tuple[int, str]] = []

        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                text = ''.join(orig_tokens[i1:i2])
                if text:
                    diffs.append((DIFF_EQUAL, text))
            elif tag == 'replace':
                del_text = ''.join(orig_tokens[i1:i2])
                ins_text = ''.join(corr_tokens[j1:j2])
                # Skip whitespace-only replacements (avoid redlining space changes)
                if del_text.strip() == '' and ins_text.strip() == '':
                    # Both are whitespace-only - just keep the original as equal
                    if del_text:
                        diffs.append((DIFF_EQUAL, del_text))
                else:
                    if del_text:
                        diffs.append((DIFF_DELETE, del_text))
                    if ins_text:
                        diffs.append((DIFF_INSERT, ins_text))
            elif tag == 'delete':
                del_text = ''.join(orig_tokens[i1:i2])
                # Skip whitespace-only deletions
                if del_text and del_text.strip() != '':
                    diffs.append((DIFF_DELETE, del_text))
                elif del_text:
                    # Whitespace-only - keep as equal to preserve spacing
                    diffs.append((DIFF_EQUAL, del_text))
            elif tag == 'insert':
                ins_text = ''.join(corr_tokens[j1:j2])
                # Skip whitespace-only insertions from redlining, but preserve the spacing
                if ins_text and ins_text.strip() != '':
                    diffs.append((DIFF_INSERT, ins_text))
                elif ins_text:
                    # Whitespace-only - keep as equal to preserve spacing
                    diffs.append((DIFF_EQUAL, ins_text))

        # Optional: merge adjacent same-op segments
        merged: List[Tuple[int, str]] = []
        for op, text in diffs:
            if merged and merged[-1][0] == op:
                merged[-1] = (op, merged[-1][1] + text)
            else:
                merged.append((op, text))
        return merged
    
    def process_document(self, file_path: str, ai_correction_func, output_path: Optional[str] = None, corrector=None) -> str:
        """
        Process an entire document, finding the boundary marker and applying
        corrections to all menu content.

        Args:
            file_path: Path to the input .docx file
            ai_correction_func: Function that takes original text and returns corrected text
            output_path: Optional output path. If None, appends "_Corrected" to input filename
            corrector: Optional AICorrector instance. If provided, its allergen codes
                      will be updated based on any allergen legend found in the document.

        Returns:
            Path to the output file
        """
        doc = Document(file_path)
        marker_found = False
        paragraphs_to_process = []

        # Find the boundary marker and collect menu paragraphs
        for para in doc.paragraphs:
            if marker_found:
                # If the marker was found, this is a menu paragraph
                paragraphs_to_process.append(para)
            elif self.boundary_marker in para.text:
                # We found the marker. Don't process this line.
                # Start processing the *next* paragraph.
                marker_found = True

        if not marker_found:
            print(f"Warning: Boundary marker '{self.boundary_marker}' not found in document.")
            print("Processing all paragraphs instead.")
            paragraphs_to_process = list(doc.paragraphs)

        # Detect document-specific allergen legend and configure corrector
        doc_allergen_codes = get_allergen_codes_for_document(paragraphs_to_process)
        if corrector is not None and hasattr(corrector, 'set_allergen_codes'):
            corrector.set_allergen_codes(doc_allergen_codes)
        
        # Check if this is a prix fixe menu and add course numbers if needed
        if self.is_prix_fixe_menu(paragraphs_to_process):
            courses_added = self.add_course_numbers(doc, paragraphs_to_process)
            if courses_added > 0:
                print(f"  Added {courses_added} course numbers")
        
        # Process each menu paragraph with AI corrections
        modified_count = 0
        for para in paragraphs_to_process:
            if self.process_paragraph(para, ai_correction_func):
                modified_count += 1
        
        print(f"Processed {len(paragraphs_to_process)} paragraphs, modified {modified_count}")
        
        # Determine output path
        if output_path is None:
            path_obj = Path(file_path)
            output_path = str(path_obj.parent / f"{path_obj.stem}_Corrected{path_obj.suffix}")
        
        # Save the document
        doc.save(output_path)
        print(f"Saved corrected document to: {output_path}")
        
        return output_path


def simple_correction_example(text: str) -> str:
    """
    Example correction function for testing.
    Replace this with actual AI integration.
    """
    # Simple example: fix common typos
    corrections = {
        'avacado': 'avocado',
        'tomatoe': 'tomato',
        'cheeze': 'cheese',
        'recieve': 'receive',
    }
    
    corrected = text
    for wrong, right in corrections.items():
        corrected = corrected.replace(wrong, right)
    
    return corrected


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python menu_redliner.py <input_file.docx> [output_file.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Create redliner and process document
    redliner = MenuRedliner()
    result = redliner.process_document(
        input_file,
        simple_correction_example,
        output_file
    )
    
    print(f"\nSuccess! Corrected document saved to: {result}")

