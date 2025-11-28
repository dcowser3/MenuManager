"""
Menu Document Redliner
======================
This module processes Word documents containing menu content and applies
AI-generated corrections with tracked changes (strikethrough for deletions,
highlight for additions) while preserving all original formatting.

The document is expected to have a boundary marker that separates the
template content from the menu content. Only content after the boundary
marker is processed.

Special handling for Prix Fixe menus:
- Detects prix fixe/tasting menus by keywords
- Automatically adds course numbers (1, 2, 3, 4...) before section headers
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import diff_match_patch as dmp_module
from typing import List, Tuple, Optional
import re
import difflib
import os
from pathlib import Path

# The text that separates the template from the menu
BOUNDARY_MARKER = "Please drop the menu content below on page 2."

# Keywords that indicate a prix fixe / tasting menu
PRIX_FIXE_KEYWORDS = [
    "prix fixe", "pre-fix", "prefix", "prix-fixe",
    "tasting menu", "tasting experience",
    "course menu", "multi-course", "multicourse",
    "degustation", "chef's menu", "chef's table"
]

# Patterns for course section headers (creative names with subtitles)
# Examples: "The Spark – "El Primer Encuentro"", "The Connection – "El Abrazo""
# Note: Use explicit Unicode escapes to ensure proper matching
COURSE_HEADER_PATTERNS = [
    # Pattern: "The X – "Spanish Subtitle"" (using en-dash U+2013 and curly quotes U+201C/U+201D)
    r'^The\s+\w+\s*[\u2013\u2014-]\s*[\u201c\u201d"][^\u201c\u201d"]+[\u201c\u201d"]',
    # Pattern: "Course Name – Subtitle" with curly or straight quotes
    r'^[A-Z][a-z]+(?:\s+[A-Z]?[a-z]+)*\s*[\u2013\u2014-]\s*[\u201c\u201d"][^\u201c\u201d"]+[\u201c\u201d"]',
    # Pattern: "COURSE X" or "Course X" followed by name
    r'^(?:COURSE|Course)\s+(?:One|Two|Three|Four|Five|Six|Seven|Eight|\d+)',
]


class MenuRedliner:
    """
    Handles the processing of menu documents with AI corrections
    and formatted diff tracking.
    """
    
    def __init__(self, boundary_marker: str = BOUNDARY_MARKER):
        """
        Initialize the MenuRedliner.
        
        Args:
            boundary_marker: The text that marks the start of menu content
        """
        self.boundary_marker = boundary_marker
        self.dmp = dmp_module.diff_match_patch()
    
    def is_prix_fixe_menu(self, paragraphs: List) -> bool:
        """
        Detect if the menu is a prix fixe / tasting menu by scanning
        for keywords in the first several paragraphs.
        
        Args:
            paragraphs: List of paragraph objects to scan
            
        Returns:
            True if this appears to be a prix fixe menu
        """
        # Check the first 10 paragraphs for prix fixe keywords
        text_to_check = " ".join(p.text.lower() for p in paragraphs[:10])
        
        for keyword in PRIX_FIXE_KEYWORDS:
            if keyword in text_to_check:
                print(f"  Detected prix fixe menu (keyword: '{keyword}')")
                return True
        
        return False
    
    def is_course_header(self, text: str) -> bool:
        """
        Check if a paragraph text matches a course section header pattern.
        
        Args:
            text: The paragraph text to check
            
        Returns:
            True if this looks like a course section header
        """
        text = text.strip()
        if not text:
            return False
        
        for pattern in COURSE_HEADER_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def add_course_numbers(self, doc, paragraphs_to_process: List) -> int:
        """
        Add course numbers before section headers in a prix fixe menu.
        Inserts new paragraphs with just the course number.
        
        Args:
            doc: The Document object
            paragraphs_to_process: List of menu paragraphs
            
        Returns:
            Number of course numbers added
        """
        course_number = 0
        paragraphs_with_numbers = []
        
        # First pass: identify which paragraphs are course headers
        for para in paragraphs_to_process:
            if self.is_course_header(para.text):
                course_number += 1
                paragraphs_with_numbers.append((para, course_number))
        
        if course_number == 0:
            print("  No course headers found to number")
            return 0
        
        print(f"  Found {course_number} course sections to number")
        
        # Second pass: insert course numbers
        # We need to insert from bottom to top to preserve indices
        for para, num in reversed(paragraphs_with_numbers):
            # Find the paragraph's position in the document body
            body = doc._body._body
            para_element = para._element
            
            # Create a new paragraph for the course number
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Create new paragraph element
            new_p = OxmlElement('w:p')
            
            # Create paragraph properties for centering (matching the style)
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            new_p.append(pPr)
            
            # Create run with the number
            run = OxmlElement('w:r')
            
            # Run properties (bold, same font)
            rPr = OxmlElement('w:rPr')
            bold = OxmlElement('w:b')
            rPr.append(bold)
            run.append(rPr)
            
            # Text content
            text_elem = OxmlElement('w:t')
            text_elem.text = str(num)
            run.append(text_elem)
            
            new_p.append(run)
            
            # Insert the new paragraph before the course header
            para_element.addprevious(new_p)
            
            # Also add yellow highlight to mark it as an addition
            from docx.oxml.ns import nsmap
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'yellow')
            rPr.append(highlight)
        
        return course_number
    
    def find_run_at_index(self, original_runs: List, original_text_len: int, target_index: int):
        """
        Helper function to find which original run (and its style)
        corresponds to a character index in the paragraph's plain text.
        
        Args:
            original_runs: List of runs from the original paragraph
            original_text_len: Total length of the original text
            target_index: Character index to find
            
        Returns:
            The run that contains the character at target_index
        """
        cumulative_len = 0
        for run in original_runs:
            run_len = len(run.text)
            if cumulative_len + run_len > target_index:
                return run
            cumulative_len += run_len
        # Fallback to the last run if index is at the very end
        return original_runs[-1] if original_runs else None
    
    def apply_formatted_diffs(self, para, diffs: List[Tuple[int, str]]):
        """
        Rebuilds a paragraph run-by-run based on diffs,
        preserving original styles and applying diff formatting.
        
        Args:
            para: The paragraph object to modify
            diffs: List of (operation, text) tuples from diff_match_patch
        """
        # 1. Save original runs and build a plain-text map
        original_runs = list(para.runs)
        original_full_text = "".join(r.text for r in original_runs)
        
        if not original_runs:
            # If there are no runs, just create new ones
            for op, text in diffs:
                if not text:
                    continue
                new_run = para.add_run(text)
                if op == dmp_module.diff_match_patch.DIFF_DELETE:
                    new_run.font.strike = True
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif op == dmp_module.diff_match_patch.DIFF_INSERT:
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return
        
        # 2. Clear the paragraph's existing content
        # We need to remove all runs while keeping the paragraph
        for run in para.runs:
            run._element.getparent().remove(run._element)
        
        current_text_index = 0
        
        # 3. Iterate through diffs
        for op, text in diffs:
            if not text:
                continue
            
            # 4. Find the style from the original run
            # This gets the run (and its style) that was at this position
            style_run = self.find_run_at_index(
                original_runs, 
                len(original_full_text), 
                current_text_index
            )
            
            # 5. Add the new run with the diff text
            new_run = para.add_run(text)
            
            # 6. Copy the original style (font, bold, italic, etc.)
            if style_run:
                # Copy font properties
                if style_run.font.name:
                    new_run.font.name = style_run.font.name
                if style_run.font.size:
                    new_run.font.size = style_run.font.size
                if style_run.font.bold is not None:
                    new_run.font.bold = style_run.font.bold
                if style_run.font.italic is not None:
                    new_run.font.italic = style_run.font.italic
                if style_run.font.underline is not None:
                    new_run.font.underline = style_run.font.underline
                if style_run.font.color.rgb:
                    # Only copy color if we're not going to override it
                    if op != dmp_module.diff_match_patch.DIFF_DELETE:
                        new_run.font.color.rgb = style_run.font.color.rgb
            
            # 7. Apply diff-specific formatting
            if op == dmp_module.diff_match_patch.DIFF_DELETE:
                # Red strikethrough for deletions
                new_run.font.strike = True
                new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red
                
            elif op == dmp_module.diff_match_patch.DIFF_INSERT:
                # Yellow highlight for additions
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
            # (op == DIFF_EQUAL needs no extra formatting)
            
            # 8. Advance the text index *only* if the text was
            # part of the original string (i.e., not an insertion)
            if op != dmp_module.diff_match_patch.DIFF_INSERT:
                current_text_index += len(text)
    
    def process_paragraph(self, para, ai_correction_func) -> bool:
        """
        Process a single paragraph with AI correction and diff application.
        
        Args:
            para: The paragraph object to process
            ai_correction_func: Function that takes original text and returns corrected text
            
        Returns:
            True if the paragraph was modified, False otherwise
        """
        original_text = para.text
        if not original_text.strip():
            # Skip empty paragraphs
            return False
        
        # Get AI correction
        corrected_text = ai_correction_func(original_text)
        
        # If no changes, skip processing
        if original_text == corrected_text:
            return False
        
        # Compute the difference using WORD-LEVEL diffs so replacements
        # appear as whole-word changes rather than single letters.
        diffs = self._word_level_diffs(original_text, corrected_text)
        
        # Apply the formatted diffs
        self.apply_formatted_diffs(para, diffs)
        
        return True

    def _tokenize(self, text: str) -> List[str]:
        """Split text into tokens: words, whitespace, and punctuation.
        This ensures diffs happen on word boundaries.
        """
        return re.findall(r"\w+|\s+|[^\w\s]", text)

    def _word_level_diffs(self, original_text: str, corrected_text: str) -> List[Tuple[int, str]]:
        """
        Produce diffs at word/token level using difflib, then map to
        diff_match_patch-style tuples. This forces whole-word replacements.
        """
        DIFF_EQUAL = dmp_module.diff_match_patch.DIFF_EQUAL
        DIFF_DELETE = dmp_module.diff_match_patch.DIFF_DELETE
        DIFF_INSERT = dmp_module.diff_match_patch.DIFF_INSERT

        orig_tokens = self._tokenize(original_text)
        corr_tokens = self._tokenize(corrected_text)

        sm = difflib.SequenceMatcher(a=orig_tokens, b=corr_tokens)
        diffs: List[Tuple[int, str]] = []

        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                text = ''.join(orig_tokens[i1:i2])
                if text:
                    diffs.append((DIFF_EQUAL, text))
            elif tag == 'replace':
                del_text = ''.join(orig_tokens[i1:i2])
                ins_text = ''.join(corr_tokens[j1:j2])
                if del_text:
                    diffs.append((DIFF_DELETE, del_text))
                if ins_text:
                    diffs.append((DIFF_INSERT, ins_text))
            elif tag == 'delete':
                del_text = ''.join(orig_tokens[i1:i2])
                if del_text:
                    diffs.append((DIFF_DELETE, del_text))
            elif tag == 'insert':
                ins_text = ''.join(corr_tokens[j1:j2])
                if ins_text:
                    diffs.append((DIFF_INSERT, ins_text))

        # Optional: merge adjacent same-op segments
        merged: List[Tuple[int, str]] = []
        for op, text in diffs:
            if merged and merged[-1][0] == op:
                merged[-1] = (op, merged[-1][1] + text)
            else:
                merged.append((op, text))
        return merged
    
    def process_document(self, file_path: str, ai_correction_func, output_path: Optional[str] = None) -> str:
        """
        Process an entire document, finding the boundary marker and applying
        corrections to all menu content.
        
        Args:
            file_path: Path to the input .docx file
            ai_correction_func: Function that takes original text and returns corrected text
            output_path: Optional output path. If None, appends "_Corrected" to input filename
            
        Returns:
            Path to the output file
        """
        doc = Document(file_path)
        marker_found = False
        paragraphs_to_process = []
        
        # Find the boundary marker and collect menu paragraphs
        for para in doc.paragraphs:
            if marker_found:
                # If the marker was found, this is a menu paragraph
                paragraphs_to_process.append(para)
            elif self.boundary_marker in para.text:
                # We found the marker. Don't process this line.
                # Start processing the *next* paragraph.
                marker_found = True
        
        if not marker_found:
            print(f"Warning: Boundary marker '{self.boundary_marker}' not found in document.")
            print("Processing all paragraphs instead.")
            paragraphs_to_process = list(doc.paragraphs)
        
        # Check if this is a prix fixe menu and add course numbers if needed
        if self.is_prix_fixe_menu(paragraphs_to_process):
            courses_added = self.add_course_numbers(doc, paragraphs_to_process)
            if courses_added > 0:
                print(f"  Added {courses_added} course numbers")
        
        # Process each menu paragraph with AI corrections
        modified_count = 0
        for para in paragraphs_to_process:
            if self.process_paragraph(para, ai_correction_func):
                modified_count += 1
        
        print(f"Processed {len(paragraphs_to_process)} paragraphs, modified {modified_count}")
        
        # Determine output path
        if output_path is None:
            path_obj = Path(file_path)
            output_path = str(path_obj.parent / f"{path_obj.stem}_Corrected{path_obj.suffix}")
        
        # Save the document
        doc.save(output_path)
        print(f"Saved corrected document to: {output_path}")
        
        return output_path


def simple_correction_example(text: str) -> str:
    """
    Example correction function for testing.
    Replace this with actual AI integration.
    """
    # Simple example: fix common typos
    corrections = {
        'avacado': 'avocado',
        'tomatoe': 'tomato',
        'cheeze': 'cheese',
        'recieve': 'receive',
    }
    
    corrected = text
    for wrong, right in corrections.items():
        corrected = corrected.replace(wrong, right)
    
    return corrected


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python menu_redliner.py <input_file.docx> [output_file.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Create redliner and process document
    redliner = MenuRedliner()
    result = redliner.process_document(
        input_file,
        simple_correction_example,
        output_file
    )
    
    print(f"\nSuccess! Corrected document saved to: {result}")

