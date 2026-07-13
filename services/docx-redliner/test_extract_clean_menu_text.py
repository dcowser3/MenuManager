#!/usr/bin/env python3
"""
Tests for extract_clean_menu_text.py — focused on track-change and
strikethrough handling that caused the "duplicate words" bug.

Run:  python -m pytest test_extract_clean_menu_text.py -v
"""

import os
import tempfile

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from lxml import etree
from docx.oxml.ns import qn, nsmap

from extract_clean_menu_text import (
    all_runs_in_paragraph,
    extract_texts,
    paragraph_clean_html,
    paragraph_clean_text,
    paragraph_unapproved_html,
    paragraph_unapproved_text,
    run_is_in_deleted_change,
    run_is_in_inserted_change,
    run_is_struck_through,
)

# ── Helpers ──────────────────────────────────────────────────────────────────

def _save_and_extract(doc, mode="clean"):
    """Save a Document to a temp file, extract, and return the result dict."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        try:
            return extract_texts(f.name, mode=mode)
        finally:
            os.unlink(f.name)


def _inject_tracked_deletion(paragraph, deleted_text):
    """Inject a w:del element with a run containing deleted_text into a paragraph.
    This simulates Word tracked-change deletions at the XML level."""
    p_elem = paragraph._p
    del_elem = etree.SubElement(p_elem, qn("w:del"))
    del_elem.set(qn("w:id"), "100")
    del_elem.set(qn("w:author"), "Test")
    r_elem = etree.SubElement(del_elem, qn("w:r"))
    dt_elem = etree.SubElement(r_elem, qn("w:delText"))
    dt_elem.text = deleted_text
    dt_elem.set(qn("xml:space"), "preserve")


def _inject_tracked_insertion(paragraph, inserted_text, bold=False):
    """Inject a w:ins element with a run containing inserted_text."""
    p_elem = paragraph._p
    ins_elem = etree.SubElement(p_elem, qn("w:ins"))
    ins_elem.set(qn("w:id"), "101")
    ins_elem.set(qn("w:author"), "Test")
    r_elem = etree.SubElement(ins_elem, qn("w:r"))
    if bold:
        rpr = etree.SubElement(r_elem, qn("w:rPr"))
        etree.SubElement(rpr, qn("w:b"))
    t_elem = etree.SubElement(r_elem, qn("w:t"))
    t_elem.text = inserted_text
    t_elem.set(qn("xml:space"), "preserve")


def _inject_move_from(paragraph, text):
    """Inject a w:moveFrom element (treated as deletion)."""
    p_elem = paragraph._p
    mf_elem = etree.SubElement(p_elem, qn("w:moveFrom"))
    mf_elem.set(qn("w:id"), "102")
    r_elem = etree.SubElement(mf_elem, qn("w:r"))
    dt_elem = etree.SubElement(r_elem, qn("w:delText"))
    dt_elem.text = text
    dt_elem.set(qn("xml:space"), "preserve")


def _inject_move_to(paragraph, text):
    """Inject a w:moveTo element (treated as insertion)."""
    p_elem = paragraph._p
    mt_elem = etree.SubElement(p_elem, qn("w:moveTo"))
    mt_elem.set(qn("w:id"), "103")
    r_elem = etree.SubElement(mt_elem, qn("w:r"))
    t_elem = etree.SubElement(r_elem, qn("w:t"))
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")


# ── Tests: Manual strikethrough (single) ────────────────────────────────────

class TestManualStrikethrough:
    def test_single_strikethrough_removed_in_clean(self):
        """Single strikethrough text should be stripped in clean mode."""
        doc = Document()
        para = doc.add_paragraph()
        run_del = para.add_run("old description ")
        run_del.font.strike = True
        run_del.font.color.rgb = RGBColor(0xFF, 0, 0)
        para.add_run("new description")

        text = paragraph_clean_text(para)
        assert text == "new description"

    def test_single_strikethrough_kept_in_unapproved(self):
        """Unapproved mode should keep struck-through text but mark it."""
        doc = Document()
        para = doc.add_paragraph()
        run_del = para.add_run("old ")
        run_del.font.strike = True
        para.add_run("new")

        text = paragraph_unapproved_text(para)
        assert text == "old new"

        html = paragraph_unapproved_html(para)
        assert 'class="existing-del"' in html
        assert "old " in html
        assert "new" in html


# ── Tests: Double strikethrough ─────────────────────────────────────────────

class TestDoubleStrikethrough:
    def test_double_strikethrough_removed_in_clean(self):
        """Double strikethrough (w:dstrike) should also be stripped in clean mode."""
        doc = Document()
        para = doc.add_paragraph()
        run_del = para.add_run("old description ")
        run_del.font.double_strike = True
        para.add_run("new description")

        text = paragraph_clean_text(para)
        assert text == "new description"

    def test_double_strikethrough_marked_in_unapproved(self):
        doc = Document()
        para = doc.add_paragraph()
        run_del = para.add_run("old ")
        run_del.font.double_strike = True
        para.add_run("new")

        html = paragraph_unapproved_html(para)
        assert 'class="existing-del"' in html
        assert "old " in html


# ── Tests: Word tracked changes (w:del / w:ins) ────────────────────────────

class TestTrackedChanges:
    def test_tracked_deletion_removed_in_clean(self):
        """Runs inside w:del should be removed in clean mode."""
        doc = Document()
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "old text ")
        para.add_run("kept text")

        text = paragraph_clean_text(para)
        assert "old text" not in text
        assert text == "kept text"

    def test_tracked_insertion_kept_in_clean(self):
        """Runs inside w:ins should be KEPT in clean mode."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("base ")
        _inject_tracked_insertion(para, "inserted")

        text = paragraph_clean_text(para)
        assert text == "base inserted"

    def test_tracked_change_pair_no_duplicates(self):
        """The core bug: a w:del + w:ins pair should not produce duplicate words."""
        doc = Document()
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "pineapple-habanero sauce ")
        _inject_tracked_insertion(para, "pineapple habanero sauce ")
        para.add_run("with chips")

        text = paragraph_clean_text(para)
        assert text == "pineapple habanero sauce with chips"
        assert "pineapple-habanero" not in text

    def test_all_tracked_changes_paragraph(self):
        """A paragraph with ONLY tracked changes (no direct runs) should work."""
        doc = Document()
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "remove this")
        _inject_tracked_insertion(para, "keep this")

        text = paragraph_clean_text(para)
        assert text == "keep this"

    def test_tracked_insertion_in_html(self):
        """Tracked insertions should appear in clean HTML output."""
        doc = Document()
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "old")
        _inject_tracked_insertion(para, "new", bold=True)

        html = paragraph_clean_html(para)
        assert "old" not in html
        assert "<strong>new</strong>" in html

    def test_tracked_changes_unapproved_marks_insertions(self):
        """Unapproved mode should mark tracked insertions.
        Note: w:del runs use w:delText (not w:t), so run.text is empty for them —
        they're naturally invisible via python-docx. Manual strikethrough (w:strike)
        uses regular w:t and IS visible in unapproved mode."""
        doc = Document()
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "deleted ")
        _inject_tracked_insertion(para, "inserted")

        html = paragraph_unapproved_html(para)
        assert 'class="existing-ins"' in html
        assert "inserted" in html


# ── Tests: w:moveFrom / w:moveTo ────────────────────────────────────────────

class TestMoveTracking:
    def test_move_from_removed_in_clean(self):
        doc = Document()
        para = doc.add_paragraph()
        _inject_move_from(para, "moved away ")
        para.add_run("remaining")

        text = paragraph_clean_text(para)
        assert text == "remaining"

    def test_move_to_kept_in_clean(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("prefix ")
        _inject_move_to(para, "moved here")

        text = paragraph_clean_text(para)
        assert text == "prefix moved here"


# ── Tests: all_runs_in_paragraph ────────────────────────────────────────────

class TestAllRunsInParagraph:
    def test_includes_nested_runs(self):
        """Should find runs inside w:del and w:ins, not just direct children."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("direct")
        _inject_tracked_deletion(para, "in-del")
        _inject_tracked_insertion(para, "in-ins")

        runs = all_runs_in_paragraph(para)
        texts = [r.text for r in runs if r.text]
        assert "direct" in texts
        # Runs inside tracked changes should be found
        assert len(runs) >= 3

    def test_direct_only_paragraph(self):
        """For a paragraph with only direct runs, should behave same as paragraph.runs."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("hello ")
        para.add_run("world")

        runs = all_runs_in_paragraph(para)
        assert len(runs) == 2
        assert runs[0].text == "hello "
        assert runs[1].text == "world"


# ── Tests: Full extraction pipeline ─────────────────────────────────────────

class TestFullExtraction:
    def test_clean_extraction_collapses_multispace_runs_in_text_and_html(self):
        """Tracked-change acceptance can leave spaces split across formatted runs."""
        doc = Document()
        doc.add_paragraph("MENU")
        para = doc.add_paragraph()
        first = para.add_run("Tulum")
        first.bold = True
        para.add_run("  ")
        second = para.add_run("Breeze")
        second.bold = True
        para.add_run("   21")

        result = _save_and_extract(doc, mode="clean")

        assert result["cleaned_menu_content"] == "Tulum Breeze 21"
        assert result["cleaned_menu_html"] == "<p><strong>Tulum</strong> <strong>Breeze</strong> 21</p>"
        assert "  " not in result["cleaned_menu_content"]
        assert "  " not in result["cleaned_menu_html"]

    def test_clean_extraction_strips_strikethrough(self):
        """End-to-end: a DOCX with strikethrough + replacement should extract cleanly."""
        doc = Document()
        # Add boundary
        doc.add_paragraph("MENU")
        # Add content with redlines
        para = doc.add_paragraph()
        run_del = para.add_run("grilled chicken ")
        run_del.font.strike = True
        run_del.font.color.rgb = RGBColor(0xFF, 0, 0)
        para.add_run("roasted chicken")

        result = _save_and_extract(doc, mode="clean")
        assert "grilled chicken" not in result["cleaned_menu_content"]
        assert "roasted chicken" in result["cleaned_menu_content"]

    def test_clean_extraction_strips_tracked_deletions(self):
        """End-to-end: tracked w:del should be stripped from cleaned output."""
        doc = Document()
        doc.add_paragraph("MENU")
        para = doc.add_paragraph()
        _inject_tracked_deletion(para, "old item ")
        _inject_tracked_insertion(para, "new item")

        result = _save_and_extract(doc, mode="clean")
        assert "old item" not in result["cleaned_menu_content"]
        assert "new item" in result["cleaned_menu_content"]

    def test_unapproved_extraction_preserves_all(self):
        """End-to-end: unapproved mode keeps deletions in visible_text."""
        doc = Document()
        doc.add_paragraph("MENU")
        para = doc.add_paragraph()
        run_del = para.add_run("struck ")
        run_del.font.strike = True
        para.add_run("kept")

        result = _save_and_extract(doc, mode="unapproved")
        assert "struck " in result["visible_text"]
        assert "kept" in result["visible_text"]
        assert 'existing-del' in result["unapproved_html"]

    def test_no_duplicate_words_in_clean(self):
        """Regression: clean extraction must not produce duplicate words from redlines."""
        doc = Document()
        doc.add_paragraph("MENU")
        para = doc.add_paragraph()
        run_del = para.add_run("seared ahi tuna ")
        run_del.font.strike = True
        run_del.font.color.rgb = RGBColor(0xFF, 0, 0)
        run_ins = para.add_run("pan-seared ahi tuna")
        run_ins.font.highlight_color = WD_COLOR_INDEX.YELLOW

        result = _save_and_extract(doc, mode="clean")
        cleaned = result["cleaned_menu_content"]
        assert cleaned.strip() == "pan-seared ahi tuna"
        # Verify no duplicate "ahi tuna"
        assert cleaned.count("ahi tuna") == 1
