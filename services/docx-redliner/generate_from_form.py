#!/usr/bin/env python3
"""
Generate a Word document from the template and form data.
This script takes the RSH menu template and populates it with form field data
and menu content, then saves it as a new document.

Usage:
    python generate_from_form.py <template_path> <form_data_json> <output_path>
"""

import sys
import json
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX


_WHITESPACE_RE = re.compile(r'\s+')
_PAREN_LEGEND_PAIR_RE = re.compile(
    r'\(\s*([A-Za-z]{1,3})\s*\)\s*([A-Za-z][A-Za-z\s/&-]*?)(?=\s*\(\s*[A-Za-z]{1,3}\s*\)|$)'
)
_PAREN_LEGEND_SPLIT_RE = re.compile(
    r'\b(?:ALL\s+PRICES|WE\s+WELCOME|CONSUMPTION\s+OF\s+RAW|CONSUMING\s+RAW|FOODBORNE\s+ILLNESS)\b',
    re.IGNORECASE,
)
_LEGEND_KEYWORD_RE = re.compile(
    r'(allergen|gluten|dairy|fish|nuts?|egg|vegan|vegetarian|crustacean|soy|sesame|celery|mustard|shellfish|sulphites?|lupin)',
    re.IGNORECASE,
)
MENU_BODY_SIDE_INDENT = Inches(1.0)


def _normalize_whitespace(value):
    return _WHITESPACE_RE.sub(' ', value or '').strip()


def is_managed_footer_line(text):
    """Return True if a menu line is an allergen legend / price footer / welcome
    blurb / raw-consumption notice. These get appended by the generator from
    structured form fields, so any inline copy in the editor body would
    duplicate. Mirrors the detection in services/dashboard/index.ts."""
    normalized = _normalize_whitespace(text)
    if not normalized:
        return False

    lower = normalized.lower()

    # "ALL PRICES ARE IN AED..." style price footer.
    if lower.startswith('all prices'):
        return True
    # "We welcome enquiries..." style allergy disclaimer.
    if lower.startswith('we welcome enquiries'):
        return True
    # Raw / undercooked consumption notice.
    if 'raw or undercooked' in lower and 'foodborne illness' in lower:
        return True
    # "Allergen Key" header line.
    if re.match(r'^allergen\s+key(?:\s+\(optional\))?$', lower):
        return True

    # Pipe-separated legend like "G gluten | D dairy | ...".
    if '|' in normalized:
        parts = [p.strip() for p in normalized.split('|') if p.strip()]
        if len(parts) >= 3:
            code_parts = [p for p in parts if re.match(r'^\*?[A-Z]{1,3}\s+.+', p)]
            if len(code_parts) >= max(2, len(parts) * 6 // 10):
                return True

    # Parenthesized legend like "(C) CELERY (D) DAIRY ...".
    if '(' in normalized and ')' in normalized:
        footer_body = _PAREN_LEGEND_SPLIT_RE.split(normalized, maxsplit=1)[0]
        pairs = _PAREN_LEGEND_PAIR_RE.findall(footer_body)
        if len(pairs) >= 4:
            hits = sum(1 for _, label in pairs if _LEGEND_KEYWORD_RE.search(label))
            if hits >= 2:
                return True

    return False


class MenuHTMLParser(HTMLParser):
    """Parse HTML content from Quill editor and convert to structured data."""

    def __init__(self):
        super().__init__()
        self.lines = []
        self.current_line = []
        self.block_stack = []
        self.format_stack = [{
            'bold': False,
            'italic': False,
            'underline': False,
            'strike': False,
            'highlight': False
        }]

    def _append_current_line(self):
        self.lines.append(self.current_line)
        self.current_line = []
        if self.block_stack:
            self.block_stack[-1]['emitted'] = True

    def _append_blank_line(self):
        self.lines.append([])
        if self.block_stack:
            self.block_stack[-1]['emitted'] = True

    def handle_starttag(self, tag, attrs):
        tag = (tag or '').lower()
        if tag in ['p', 'div']:
            if self.current_line:
                self._append_current_line()
            self.block_stack.append({'had_content': False, 'emitted': False})
        elif tag == 'br':
            if self.current_line:
                self._append_current_line()
            else:
                self._append_blank_line()
            return

        current = dict(self.format_stack[-1])
        attr_map = dict(attrs or [])
        class_attr = attr_map.get('class', '')
        class_names = set(class_attr.split()) if class_attr else set()

        if tag in ['strong', 'b']:
            current['bold'] = True
        elif tag in ['em', 'i']:
            current['italic'] = True
        elif tag == 'u':
            current['underline'] = True
        elif tag in ['s', 'del', 'strike']:
            current['strike'] = True
        elif tag == 'span':
            if 'persistent-del' in class_names or 'existing-del' in class_names:
                current['strike'] = True
            if 'persistent-ins' in class_names or 'existing-ins' in class_names:
                current['highlight'] = True

        self.format_stack.append(current)

    def handle_startendtag(self, tag, attrs):
        if (tag or '').lower() == 'br':
            self.handle_starttag(tag, attrs)
            return
        self.handle_starttag(tag, attrs)
        self.handle_endtag(tag)

    def handle_endtag(self, tag):
        tag = (tag or '').lower()
        if tag in ['p', 'div'] and self.block_stack:
            block = self.block_stack[-1]
            if self.current_line:
                self._append_current_line()
            elif not block.get('had_content') and not block.get('emitted'):
                self._append_blank_line()
            self.block_stack.pop()

        if len(self.format_stack) > 1:
            self.format_stack.pop()

    def handle_data(self, data):
        if not data:
            return

        current = self.format_stack[-1]

        # Newlines inside text behave like <br> — flush the line.
        # Otherwise python-docx would render them as <w:br/> soft breaks inside
        # one paragraph, and consecutive \n become visible blank lines in Word.
        chunks = data.split('\n')
        for i, chunk in enumerate(chunks):
            if i > 0 and self.current_line:
                self.lines.append(self.current_line)
                self.current_line = []

            if not chunk:
                continue

            # Drop whitespace-only data when no content has started on this line
            # (whitespace between tags at paragraph start). Preserve it mid-line
            # so " " between adjacent <span>s isn't lost (e.g. "Carne Asada").
            if not chunk.strip() and not self.current_line:
                continue

            self.current_line.append({
                'text': chunk,
                'bold': current['bold'],
                'italic': current['italic'],
                'underline': current['underline'],
                'strike': current['strike'],
                'highlight': current['highlight']
            })
            if self.block_stack and (chunk.strip() or self.current_line):
                self.block_stack[-1]['had_content'] = True

    def get_lines(self):
        """Get parsed lines with formatting."""
        if self.current_line:
            self.lines.append(self.current_line)
        return self.lines


def populate_template(template_path: str, form_data: dict, output_path: str):
    """
    Populate the Word template with form data.

    Args:
        template_path: Path to the RSH menu template .docx file
        form_data: Dictionary containing form fields (projectName, property, etc.)
        output_path: Path where the populated document should be saved
    """
    print(f"Loading template from: {template_path}")
    doc = Document(template_path)

    # Map form fields to table cells
    # The template has a table with form fields on page 1
    field_mapping = {
        'PROJECT NAME': form_data.get('projectName', ''),
        'PROPERTY': form_data.get('property', ''),
        'SIZE (PIXELS = WEB) OR (INCHES = PRINT)': form_data.get('size', ''),
        'ORIENTATION (PORTRAIT OR LANDSCAPE)': form_data.get('orientation', ''),
        'DATE NEEDED': form_data.get('dateNeeded', '')
    }

    print("Populating form fields in template...")

    # Find and populate the table (should be first table in document)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                field_name = cells[0].text.strip()
                if field_name in field_mapping:
                    # Clear existing content and add new value
                    cells[1].text = field_mapping[field_name]
                    print(f"  ✓ Set '{field_name}' = '{field_mapping[field_name]}'")

    # Find the boundary marker paragraph
    # Try different markers for food vs beverage templates
    boundary_markers = [
        "Please drop the menu content below on page 2",  # Food template
        "MENU"  # Beverage template (just "MENU" on its own line)
    ]
    boundary_index = None

    for i, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        # Check for exact "MENU" match (beverage template)
        if para_text == "MENU":
            boundary_index = i
            print(f"Found 'MENU' marker at paragraph {i}")
            break
        # Check for the food template marker
        elif boundary_markers[0] in para_text:
            boundary_index = i
            print(f"Found boundary marker at paragraph {i}")
            break

    if boundary_index is None:
        print("WARNING: Could not find boundary marker in template")
        boundary_index = len(doc.paragraphs) - 1

    # Remove any existing content after the boundary marker
    # (template might have placeholder text)
    paragraphs_to_remove = []
    for i in range(boundary_index + 1, len(doc.paragraphs)):
        paragraphs_to_remove.append(doc.paragraphs[i])

    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

    print("Adding menu content...")

    def apply_menu_paragraph_style(para):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Keep generated menu paragraphs compact and predictable in Word.
        para.paragraph_format.left_indent = MENU_BODY_SIDE_INDENT
        para.paragraph_format.right_indent = MENU_BODY_SIDE_INDENT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15

    # Check if we have HTML content to preserve formatting
    menu_content_html = form_data.get('menuContentHtml', '')
    menu_content_text = form_data.get('menuContent', '')
    allergens_text = (form_data.get('allergens', '') or '').strip()
    footer_text = (form_data.get('footerText', '') or '').strip()
    should_add_raw_notice = bool(form_data.get('shouldAddRawNotice', False))

    # Force menu body to start on a fresh page regardless of template flow.
    page_break_paragraph = doc.add_paragraph()
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    if menu_content_html:
        # Parse HTML content to preserve formatting
        parser = MenuHTMLParser()
        parser.feed(menu_content_html)
        lines = parser.get_lines()

        rendered = 0
        skipped_footer = 0
        for line_parts in lines:
            line_text = ''.join(part.get('text', '') for part in line_parts)
            if is_managed_footer_line(line_text):
                skipped_footer += 1
                continue

            para = doc.add_paragraph()
            apply_menu_paragraph_style(para)

            for part in line_parts:
                run = para.add_run(part['text'])
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.bold = part['bold']
                run.italic = part['italic']
                run.underline = part['underline']
                run.font.strike = bool(part.get('strike'))
                if part.get('strike'):
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                if part.get('highlight'):
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            rendered += 1

        print(f"Added {rendered} lines of formatted menu content (skipped {skipped_footer} inline footer lines)")
    else:
        # Fall back to plain text if no HTML
        lines = menu_content_text.split('\n')

        rendered = 0
        skipped_footer = 0
        for line in lines:
            if line.strip() and is_managed_footer_line(line):
                skipped_footer += 1
                continue

            para = doc.add_paragraph()
            apply_menu_paragraph_style(para)

            if line.strip():
                run = para.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            rendered += 1

        print(f"Added {rendered} lines of plain text menu content (skipped {skipped_footer} inline footer lines)")

    # Append allergen legend in a compact, single-line format.
    if allergens_text:
        doc.add_paragraph()

        # Support either pipe-delimited single line or multi-line key definitions,
        # then normalize to one line separated by " | ".
        allergen_lines = [line.strip() for line in allergens_text.splitlines() if line.strip()]
        if len(allergen_lines) == 1 and '|' in allergen_lines[0]:
            allergen_lines = [part.strip() for part in allergen_lines[0].split('|') if part.strip()]
        normalized_allergen_line = " | ".join(allergen_lines)

        allergen_para = doc.add_paragraph()
        apply_menu_paragraph_style(allergen_para)
        allergen_run = allergen_para.add_run(normalized_allergen_line)
        allergen_run.font.name = 'Calibri'
        allergen_run.font.size = Pt(10)

        if footer_text:
            for footer_line in [line.strip() for line in footer_text.splitlines() if line.strip()]:
                footer_para = doc.add_paragraph()
                apply_menu_paragraph_style(footer_para)
                footer_run = footer_para.add_run(footer_line)
                footer_run.font.name = 'Calibri'
                footer_run.font.size = Pt(10)

        if should_add_raw_notice:
            # Add the raw-consumption notice underneath the allergen legend when requested.
            raw_notice_para = doc.add_paragraph()
            apply_menu_paragraph_style(raw_notice_para)
            raw_notice_run = raw_notice_para.add_run(
                "*consuming raw or undercooked meats, poultry, seafood, shellfish, or eggs may increase your risk of foodborne illness."
            )
            raw_notice_run.font.name = 'Calibri'
            raw_notice_run.font.size = Pt(10)

    # Save the populated document
    print(f"Saving document to: {output_path}")
    doc.save(output_path)
    print("✓ Document generated successfully")


def main():
    if len(sys.argv) != 4:
        print("Usage: python generate_from_form.py <template_path> <form_data_json> <output_path>")
        sys.exit(1)

    template_path = sys.argv[1]
    form_data_path = sys.argv[2]
    output_path = sys.argv[3]

    # Load form data
    try:
        with open(form_data_path, 'r') as f:
            form_data = json.load(f)
        print(f"Loaded form data: {list(form_data.keys())}")
    except Exception as e:
        print(f"ERROR: Could not load form data from {form_data_path}: {e}")
        sys.exit(1)

    # Generate document
    try:
        populate_template(template_path, form_data, output_path)
    except Exception as e:
        print(f"ERROR: Failed to generate document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
