#!/usr/bin/env python3
"""Tests for project-detail extraction from populated RSH template DOCX files."""

import os
import tempfile

from docx import Document

from extract_project_details import detect_allergen_key, extract_project_details


def _save_and_extract(doc):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        try:
            return extract_project_details(f.name)
        finally:
            os.unlink(f.name)


def test_extracts_split_property_fields_from_new_template():
    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    rows = [
        ("MENU NAME", "Dinner Menu"),
        ("OUTLET NAME", "Maya"),
        ("HOTEL NAME", "Le Royal Meridien"),
        ("CITY / COUNTRY", "Dubai"),
        ("SIZE (PIXELS = WEB) OR (INCHES = PRINT)", "8.5 x 11 inches"),
        ("ORIENTATION (PORTRAIT OR LANDSCAPE)", "Portrait"),
        ("DATE NEEDED", "2026-05-20"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    doc.add_paragraph("MENU")
    doc.add_paragraph("Taco - grilled fish")

    result = _save_and_extract(doc)

    assert result["project_details"] == {
        "project_name": "Dinner Menu",
        "property": "",
        "outlet": "Maya",
        "hotel": "Le Royal Meridien",
        "city": "Dubai",
        "size": "8.5 x 11 inches",
        "orientation": "Portrait",
        "date_needed": "2026-05-20",
    }
    assert result["menu_content"] == "Taco - grilled fish"


def test_detects_parenthesized_allergen_key():
    doc = Document()
    doc.add_paragraph("(C) CELERY (D) DAIRY (E) EGGS (F) FISH (G) GLUTEN (L) LUPIN")
    doc.add_paragraph("(M) MUSTARD (P) PORK (PN) PEANUTS (S) SHELLFISH (SL) SULPHITES")
    doc.add_paragraph("(SS) SESAME (SY) SOY (TN) TREE NUTS (V) VEGETARIAN")

    assert detect_allergen_key(doc.paragraphs) == (
        "C celery | D dairy | E eggs | F fish | G gluten | L lupin | "
        "M mustard | P pork | PN peanuts | S shellfish | SL sulphites | "
        "SS sesame | SY soy | TN tree nuts | V vegetarian"
    )


def test_parenthesized_allergen_key_stops_before_footer_copy():
    doc = Document()
    doc.add_paragraph(
        "(C) CELERY (D) DAIRY (E) EGGS (F) FISH (G) GLUTEN (V) VEGETARIAN "
        "ALL PRICES ARE IN AED, INCLUSIVE OF 7% MUNICIPALITY FEES, 10% SERVICE CHARGE AND 5% VAT."
    )

    assert detect_allergen_key(doc.paragraphs) == (
        "C celery | D dairy | E eggs | F fish | G gluten | V vegetarian"
    )


def test_extract_project_details_returns_parenthesized_allergen_key():
    doc = Document()
    doc.add_paragraph("MENU")
    doc.add_paragraph("Ice Cream & Sorbets D,E,G,PN,SY,TN 35")
    doc.add_paragraph("(D) DAIRY (E) EGGS (G) GLUTEN (PN) PEANUTS (SY) SOY (TN) TREE NUTS")

    result = _save_and_extract(doc)

    assert result["allergen_key"] == "D dairy | E eggs | G gluten | PN peanuts | SY soy | TN tree nuts"


if __name__ == "__main__":
    test_extracts_split_property_fields_from_new_template()
    test_detects_parenthesized_allergen_key()
    test_parenthesized_allergen_key_stops_before_footer_copy()
    test_extract_project_details_returns_parenthesized_allergen_key()
