#!/usr/bin/env python3
"""
Format Linter for RSH Menu Documents
Checks page 2+ (after boundary marker) for:
- Center alignment
- Font family (Calibri / Calibri (Body))
- Font size (12 pt)
Outputs JSON with pass/fail and reasons.
"""
import sys
import json
from pathlib import Path
from typing import List, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

BOUNDARY_MARKER = "Please drop the menu content below on page 2."


def paragraph_text(para) -> str:
    return "".join(run.text for run in para.runs).strip()


def is_centered(para) -> bool:
    # Some documents inherit center via styles; explicit None may still render centered.
    # We consider CENTER only when alignment explicitly set to CENTER.
    return para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def check_font_rules(para) -> Dict[str, bool]:
    """Check runs for size==12pt and calibri-like font when specified."""
    issues = {"size_mismatch": False, "font_mismatch": False}
    for run in para.runs:
        text = (run.text or "").strip()
        if not text:
            continue
        # Size: if explicitly set and not 12pt (152400 EMU) -> mismatch
        if run.font.size is not None:
            # font.size is a Length (EMU). 12 pt == 152400
            try:
                if int(run.font.size) != 152400:
                    issues["size_mismatch"] = True
            except Exception:
                issues["size_mismatch"] = True
        # Font: if explicitly set and not Calibri variants -> mismatch
        if run.font.name:
            name = run.font.name.lower()
            if "calibri" not in name:
                issues["font_mismatch"] = True
    return issues


def lint(file_path: str) -> Dict:
    doc = Document(file_path)
    marker_found = False
    menu_paragraphs: List = []

    for para in doc.paragraphs:
        if marker_found:
            menu_paragraphs.append(para)
        elif BOUNDARY_MARKER in para.text:
            marker_found = True

    if not marker_found:
        # fallback: treat entire doc
        menu_paragraphs = [p for p in doc.paragraphs if paragraph_text(p)]

    # filter out empty lines
    menu_paragraphs = [p for p in menu_paragraphs if paragraph_text(p)]

    total = len(menu_paragraphs)
    centered_offenders = []
    size_offenders = []
    font_offenders = []

    for p in menu_paragraphs:
        text = paragraph_text(p)
        if not is_centered(p):
            centered_offenders.append(text[:120])
        fr = check_font_rules(p)
        if fr["size_mismatch"]:
            size_offenders.append(text[:120])
        if fr["font_mismatch"]:
            font_offenders.append(text[:120])

    # Pass criteria: majority centered (>=80%) AND no explicit size!=12 on any run AND no explicit non-calibri fonts
    centered_ratio = 0.0 if total == 0 else (total - len(centered_offenders)) / total
    pass_center = centered_ratio >= 0.8
    pass_size = len(size_offenders) == 0
    pass_font = len(font_offenders) == 0

    passed = pass_center and pass_size and pass_font

    reasons = []
    if not pass_center:
        reasons.append(f"Menu paragraphs are not centered per SOP (only {int(centered_ratio*100)}% centered).")
    if not pass_size:
        reasons.append("Detected paragraphs with explicit font size not equal to 12 pt.")
    if not pass_font:
        reasons.append("Detected paragraphs with explicit font not Calibri (Body).")

    return {
        "passed": passed,
        "checks": {
            "center_alignment": pass_center,
            "font_size_12pt": pass_size,
            "font_family_calibri": pass_font,
        },
        "totals": {
            "menu_paragraphs": total,
        },
        "samples": {
            "not_centered": centered_offenders[:5],
            "size_mismatch": size_offenders[:5],
            "font_mismatch": font_offenders[:5],
        },
        "reasons": reasons,
    }


def main():
    if len(sys.argv) < 2:
        print("Usage: format_lint.py <file.docx>")
        sys.exit(2)
    file_path = sys.argv[1]
    if not Path(file_path).exists():
        print(json.dumps({"error": f"file not found: {file_path}"}))
        sys.exit(2)
    report = lint(file_path)
    print(json.dumps(report))


if __name__ == "__main__":
    main()




